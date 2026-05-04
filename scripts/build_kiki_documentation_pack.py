"""Build refreshed Kiki simulator DOCX guides.

The generated documents are intentionally separate from the older manuals so the
previous documentation pack remains available for comparison.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "documentation"
FIG_DIR = DOC_DIR / "figures" / "game_big_picture"
ASSET_DIR = ROOT / "assets"
KIKI_IMAGE = ASSET_DIR / "kiki_mascot.png"

BLUE = "2563EB"
GREEN = "16A34A"
ORANGE = "D97706"
PURPLE = "7C3AED"
RED = "DC2626"
TEAL = "0D9488"
INK = "172033"
MUTED = "59677B"
LIGHT_BLUE = "DBEAFE"
LIGHT_GREEN = "DCFCE7"
LIGHT_ORANGE = "FEF3C7"
LIGHT_PURPLE = "EDE9FE"
LIGHT_RED = "FEE2E2"
LIGHT_TEAL = "CCFBF1"
LIGHT_GRAY = "F8FAFC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margin: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def style_doc(doc: Document, landscape: bool = False) -> None:
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
    else:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.62)
        section.bottom_margin = Inches(0.62)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    for style_name, size, color in [
        ("Title", 24, INK),
        ("Heading 1", 16, INK),
        ("Heading 2", 12.5, INK),
        ("Heading 3", 11, INK),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    styles["Caption"].font.name = "Aptos"
    styles["Caption"].font.size = Pt(8.5)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor.from_string(MUTED)


def compact_paragraph(paragraph, before: int = 0, after: int = 4, line: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_title_block(doc: Document, title: str, subtitle: str, audience: str, accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    section = doc.sections[0]
    usable_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    set_table_width(table, [usable_width])
    left = table.rows[0].cells[0]
    set_cell_shading(left, LIGHT_GRAY)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = left.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(accent)
    compact_paragraph(p, after=5)

    p = left.add_paragraph(subtitle)
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    compact_paragraph(p, after=5)

    p = left.add_paragraph(audience)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(9.5)
    compact_paragraph(p, after=0)


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE, accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150)
    p = cell.paragraphs[0]
    p.add_run(f"{label}: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)
    compact_paragraph(p, after=0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        compact_paragraph(p, after=1)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        compact_paragraph(p, after=1)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    header_fill: str = LIGHT_BLUE,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, 95)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(8.8)
        run.font.color.rgb = RGBColor.from_string(INK)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_margins(cell, 95)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.add_run(value)
            p.runs[0].font.size = Pt(8.6)
            compact_paragraph(p, after=0, line=1.0)


def add_section(doc: Document, title: str, intro: str | None = None) -> None:
    h = doc.add_heading(title, level=1)
    compact_paragraph(h, before=8, after=4)
    if intro:
        p = doc.add_paragraph(intro)
        p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        compact_paragraph(p, after=6)


def add_figure(doc: Document, path: Path, caption: str, width: float) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        compact_paragraph(p, after=2)
        cap = doc.add_paragraph(caption, style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact_paragraph(cap, after=4)


def build_instructor_manual() -> Path:
    doc = Document()
    style_doc(doc)
    add_title_block(
        doc,
        "Instructor Operating Manual",
        "Kiki Pickleball Business Simulation",
        "Audience: instructor, TA, simulation facilitator | Updated for Kiki branding, finance detail, team tracker, and big-picture guide.",
        PURPLE,
    )
    add_callout(
        doc,
        "Important",
        "Students should use the hosted Render URL during class, not localhost. Localhost only works on the computer running Streamlit.",
        LIGHT_RED,
        RED,
    )

    add_section(doc, "1. Simulator Purpose and Teaching Goals")
    add_table(
        doc,
        ["Concept", "What students practice", "Instructor debrief angle"],
        [
            ["Forecasting and S&OP", "Forecast product demand, align production/materials/cash.", "Did the plan match the market?"],
            ["Operations and SCM", "Capacity, overtime, supplier mix, QC, inventory, backlog.", "What constraint limited performance?"],
            ["Portfolio and lifecycle", "Manage Product A/B/C, retirement, replacement.", "Which products carried or dragged the team?"],
            ["NPD and technology", "Invest in pipeline projects and time launches.", "Did innovation arrive soon enough to matter?"],
            ["Finance and liquidity", "Balance cash, borrowing, debt, interest, working capital.", "Did profitable growth still create cash stress?"],
        ],
        [1.35, 2.55, 2.45],
        LIGHT_PURPLE,
    )

    add_section(doc, "2. User Roles and Login Workflow")
    add_table(
        doc,
        ["Role", "Access", "Cannot do"],
        [
            ["admin", "All pages, user management, market reports, instructor panel, all results.", "Should not be shared with students."],
            ["team_leader", "Home, market report, own decisions, safe dashboard, own tracker, guides.", "Cannot run rounds, reset data, edit other teams, or manage users."],
        ],
        [1.15, 3.0, 2.25],
        LIGHT_BLUE,
    )

    add_section(doc, "3. Pre-Class Setup Checklist")
    add_table(
        doc,
        ["Done", "Task", "Why it matters"],
        [
            ["[ ]", "Open hosted Render URL and log in as admin.", "Confirms the deployed app is alive."],
            ["[ ]", "Create or verify team leader accounts.", "Each team needs its own private submission identity."],
            ["[ ]", "Set starting cash when creating teams if using finance realism.", "Makes liquidity discussion intentional."],
            ["[ ]", "Open Public Market Report and set Round 1.", "Defines the business environment."],
            ["[ ]", "Open team submissions.", "Allows student decisions to be saved."],
            ["[ ]", "Test one team account.", "Confirms students can log in and submit."],
        ],
        [0.55, 3.1, 2.75],
        LIGHT_GREEN,
    )
    add_callout(
        doc,
        "Student-facing reminder",
        "Start with the market report, forecast before producing, make supplier mix total 100%, check cash before expanding, and click Save before the deadline.",
        LIGHT_BLUE,
        BLUE,
    )
    doc.add_page_break()

    add_section(doc, "4. Page-by-Page Instructor Guide")
    add_table(
        doc,
        ["Page", "Purpose", "Instructor action"],
        [
            ["Home", "Kiki-branded landing page and navigation overview.", "Use to orient students to the current version."],
            ["Public Market Report", "Set demand, segment shares, cost, risk, quality, and technology pressure.", "Apply a preset or manually tune the market before teams submit."],
            ["Team Decisions", "Students enter firm, product, forecast, finance, and pipeline choices.", "Use admin view only for inspection/testing; students see only their team."],
            ["Instructor Panel", "Submission status, validation, round execution, reset controls.", "Close submissions, check warnings, run round."],
            ["Admin User Management", "Create, reset, deactivate, bulk-create, and remove team accounts.", "Build the class roster and distribute IDs/passwords."],
            ["Results Dashboard", "Team rankings, product detail, forecast accuracy, liquidity, pipeline.", "Use tabs for debrief and class discussion."],
            ["Finance Detail", "Cash bridge, cost breakdown, debt and product contribution detail.", "Explain why profit and cash are different."],
            ["My Team Tracker", "Round-by-round team decisions and performance history.", "Let teams self-diagnose without exposing others' private inputs."],
            ["Game Big Picture", "Infographics showing the whole simulation logic.", "Use for introduction and concept review."],
            ["Model Formula Guide", "Transparent formulas and model logic.", "Use when students ask why results happened."],
            ["My Account", "Password change.", "Remind users to update temporary passwords."],
        ],
        [1.35, 2.55, 2.45],
        LIGHT_TEAL,
    )

    add_section(doc, "5. Public Market Report Controls")
    add_table(
        doc,
        ["Control", "Meaning", "Teaching use"],
        [
            ["Total Demand", "Total market size for the round.", "Boom/recession discussion."],
            ["Premium/Mid/Beginner Share", "Segment mix; should sum near 1.0.", "Segmentation and positioning."],
            ["Material Cost Index", "Multiplier on input/material cost.", "Cost inflation and supplier pressure."],
            ["Supply Risk", "Low/Moderate/High supply disruption pressure.", "Supplier mix and QC tradeoffs."],
            ["Quality Sensitivity", "How strongly defects hurt demand/reputation.", "Quality is valuable but not free."],
            ["Current Market Generation", "Technology generation customers expect.", "Obsolescence and innovation timing."],
            ["Technology Shift Rate", "How quickly tech expectations move.", "NPD urgency."],
            ["Premium/Mid Tech Adoption", "How much each segment rewards newer tech.", "Segment-specific innovation value."],
            ["Beginner Price Pressure", "How strongly beginner buyers punish high price.", "Volume vs margin."],
        ],
        [1.55, 2.6, 2.25],
        LIGHT_ORANGE,
    )

    add_section(doc, "6. Running Each Round")
    add_numbered(
        doc,
        [
            "Set or confirm the Public Market Report for the current round.",
            "Open team submissions in the Instructor Panel.",
            "Teams review the market report, enter decisions, inspect planning analytics, and save.",
            "Monitor validation: missing forecasts, supplier mix, infeasible plans, cash shortfalls, and launch-ready projects.",
            "Close team submissions when time expires.",
            "Click Run Round. The engine allocates demand, applies constraints, computes finance, and updates state.",
            "Use Results Dashboard, Finance Detail, My Team Tracker, and Game Big Picture for debrief.",
            "Return to Public Market Report, advance to the next round, set the next environment, and reopen submissions.",
        ],
    )
    add_callout(
        doc,
        "Warning",
        "Reset is for testing or restarting only. Factory reset removes runtime progress, decisions, results, team state, products, projects, forecasts, and accounts depending on the selected control.",
        LIGHT_RED,
        RED,
    )

    add_section(doc, "7. Debrief Prompts")
    add_bullets(
        doc,
        [
            "Which team had the best balance between profit, service, forecast accuracy, cash health, and innovation timing?",
            "Which team sold a lot but tied up too much cash in inventory, expansion, or debt?",
            "Which product generated profit and which product quietly dragged the portfolio?",
            "Did any team forecast demand well but fail because capacity, materials, or cash did not support the plan?",
            "Did a new technology launch cannibalize an older product or create a true competitive advantage?",
        ],
    )

    add_section(doc, "8. Common Troubleshooting")
    add_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Student cannot log in", "Wrong username/password or inactive account.", "Reset password or activate user in Admin User Management."],
            ["Team cannot save", "Submissions closed or validation error.", "Open submissions or fix supplier mix/required fields."],
            ["Round results look missing", "Round not run yet or no team submitted.", "Check Instructor Panel submission table, then run round."],
            ["Public Market Report still shows old round", "Next round not saved after running.", "Set the next round number and save the report."],
            ["Render page error after deploy", "Old SQLite schema or stale disk.", "Use migrations/reset tools; if testing only, factory reset the runtime data."],
        ],
        [1.7, 2.25, 2.45],
        LIGHT_RED,
    )
    add_section(doc, "9. Quick Live-Class Reference")
    add_table(
        doc,
        ["Moment", "Instructor check", "Student message"],
        [
            ["Opening", "Hosted URL works, accounts active, market report saved.", "Log in, read the market, and start from your current state."],
            ["Decision time", "Submission table, missing forecasts, supplier mix, infeasible plans.", "Forecast first, then align production, materials, capacity, and cash."],
            ["Before Run Round", "Close submissions, scan warnings, decide whether late fixes are allowed.", "Only saved decisions count."],
            ["Debrief", "Rankings, product detail, finance detail, tracker history.", "Explain your tradeoff: profit, service, forecast accuracy, cash, or innovation."],
        ],
        [1.4, 2.55, 2.45],
        LIGHT_TEAL,
    )
    add_callout(
        doc,
        "Instructor tip",
        "If students ask who is winning, ask by which measure: profit, service, forecast discipline, cash health, or innovation timing. That is the central lesson.",
        LIGHT_ORANGE,
        ORANGE,
    )
    add_table(
        doc,
        ["Dashboard area", "Use it to ask", "Typical lesson"],
        [
            ["Team Summary", "Who appears to be winning, and why?", "Rankings depend on the objective."],
            ["Forecast Accuracy", "Who planned close to actual demand?", "Good plans reduce operational surprises."],
            ["Liquidity / Debt", "Who made profit but still created cash pressure?", "Growth can fail financially."],
            ["Product Detail", "Which product made or lost money?", "Portfolio choices matter more than firm averages."],
            ["Development Pipeline", "Who invested early enough to launch?", "Innovation requires timing and discipline."],
        ],
        [1.45, 2.45, 2.5],
        LIGHT_BLUE,
    )

    out = DOC_DIR / "Kiki_Instructor_Operating_Manual_v3.docx"
    doc.save(out)
    return out


def build_model_guide() -> Path:
    doc = Document()
    style_doc(doc)
    add_title_block(
        doc,
        "Simulator Model and Formula Guide",
        "Kiki Pickleball Business Simulation",
        "Audience: instructor, TA, researcher, or advanced student who wants to understand the deterministic engine.",
        TEAL,
    )
    add_callout(
        doc,
        "Model stance",
        "The simulator is a transparent deterministic teaching model, not an AI black box. Coefficients are tunable calibration choices for classroom learning.",
        LIGHT_TEAL,
        TEAL,
    )

    add_section(doc, "1. Round Engine Sequence")
    add_numbered(
        doc,
        [
            "Load market report, team decisions, product lines, product decisions, development projects, and persistent team state.",
            "Prepare each team: installed capacity, raw material inventory, supplier mix, cash, debt, reputation, products, and projects.",
            "Update NPD progress and determine launch readiness gates.",
            "Apply launch, retirement, and replacement decisions.",
            "Cap production by shared capacity and raw material constraints.",
            "Compute product-level defect rates, good units, demand attractiveness, demand allocation, sales, backlog, inventory, and product contribution.",
            "Apply intra-team cannibalization where newer/similar products steal demand from older products.",
            "Aggregate product results to team-level revenue, cost, profit, cash, debt, liquidity, forecast accuracy, reputation, and state updates.",
        ],
    )

    add_section(doc, "2. Demand Allocation")
    add_table(
        doc,
        ["Formula / component", "Meaning"],
        [
            ["premium_demand = total_demand x premium_share", "Market report splits demand into premium, mid, and beginner segments."],
            ["attractiveness score", "Weighted score from price fit, quality, segment fit, service readiness, technology, lifecycle, reputation, and product demand fit."],
            ["product_demand = segment_demand x product_score / total_segment_score", "Each product earns a proportional share of segment demand."],
            ["cannibalization transfer", "Same-team products in the same group can transfer some demand from older to newer/stronger products."],
        ],
        [2.45, 3.95],
        LIGHT_BLUE,
    )

    add_section(doc, "3. Operations and Cost")
    add_table(
        doc,
        ["Area", "Formula idea", "Student implication"],
        [
            ["Capacity", "effective_capacity = installed_capacity + overtime", "Production cannot exceed what the firm can make."],
            ["Materials", "raw_material_available = beginning_RM + current_order_effect", "Low raw material support caps output."],
            ["Defects", "defect_rate = base + supplier pressure + utilization stress - QC effect", "QC reduces defects but costs money."],
            ["Good units", "good_units = production x (1 - defect_rate)", "Not every produced unit is sellable."],
            ["Sales", "sales = min(allocated demand, available units)", "Demand alone is not enough; inventory/service matter."],
            ["Profit", "profit = revenue - total_cost", "Costs include procurement, production, QC, holding, warranty, backlog, expansion, NPD, and interest."],
        ],
        [1.2, 2.7, 2.5],
        LIGHT_GREEN,
    )

    add_section(doc, "4. Forecast and S&OP Metrics")
    add_table(
        doc,
        ["Metric", "Formula", "Meaning"],
        [
            ["Forecast error", "actual_demand - forecast", "Positive means under-forecast; negative means over-forecast."],
            ["Absolute error", "abs(actual_demand - forecast)", "Size of miss regardless of direction."],
            ["WAPE", "sum(abs(actual - forecast)) / max(sum(actual), 1)", "Portfolio-level weighted forecast accuracy."],
            ["Forecast-production gap", "planned_production - forecast", "Shows whether production plan matches expected demand."],
        ],
        [1.4, 2.7, 2.3],
        LIGHT_BLUE,
    )

    add_section(doc, "5. Cash, Debt, and Liquidity")
    add_table(
        doc,
        ["Mechanic", "Formula / rule", "Effect"],
        [
            ["Cash bridge", "ending_cash_before_borrowing = starting_cash + revenue + planned_borrowing - costs - interest", "Shows whether operations generated or consumed cash."],
            ["Planned borrowing", "Student enters borrowing amount.", "Cash rises now, debt rises by same amount, interest applies."],
            ["Automatic borrowing", "If cash before borrowing < 0, auto-borrow the shortfall.", "Ending cash becomes zero, debt increases, liquidity stress is flagged."],
            ["Interest", "interest = debt_balance x periodic_interest_rate", "Debt reduces future profit."],
            ["Working capital", "FG inventory value + RM inventory value + debt burden factor", "Inventory and debt tie up financial flexibility."],
            ["Liquidity stress", "Flagged by low cash, high debt/revenue, high working-capital/revenue, repeated borrowing.", "Can reduce reputation and signal unsustainable growth."],
        ],
        [1.45, 2.8, 2.15],
        LIGHT_RED,
    )

    add_section(doc, "6. Product Development and Launch Readiness")
    add_table(
        doc,
        ["Gate", "How it works", "Managerial meaning"],
        [
            ["Funding progress", "Cumulative investment compared with required investment.", "More money generally speeds funding progress."],
            ["Readiness score", "Investment progress plus testing effect, reduced by higher technology complexity.", "High-tech projects need more money/testing."],
            ["Earliest launch timing", "A project cannot launch before its timing gate.", "Prevents instant launches."],
            ["Launch decision", "Team must check Launch Now when eligible.", "Readiness alone does not automatically launch a product."],
            ["Fixed settings", "Project name, segment, tech generation, slot, launch timing lock after project begins.", "Students manage execution, not endless redesign."],
        ],
        [1.35, 2.8, 2.25],
        LIGHT_PURPLE,
    )

    add_section(doc, "7. Performance Measures")
    add_table(
        doc,
        ["Measure", "Why it matters"],
        [
            ["Profit", "Short-run economic success after all costs."],
            ["Fill rate / service", "Ability to meet demand with available sellable units."],
            ["Forecast accuracy", "Planning discipline and S&OP quality."],
            ["Cash minus debt", "Financial flexibility and survival."],
            ["Reputation", "Future demand impact from service, quality, and liquidity stress."],
            ["Innovation position", "Pipeline readiness, tech generation, launch timing, and replacement decisions."],
        ],
        [2.0, 4.4],
        LIGHT_TEAL,
    )

    add_section(doc, "8. Calibration Note")
    add_callout(
        doc,
        "Instructor tip",
        "Do not score only by profit. A balanced score can include profit, service, forecast accuracy, cash health, and innovation timing so students learn tradeoffs rather than gaming one number.",
        LIGHT_ORANGE,
        ORANGE,
    )
    add_table(
        doc,
        ["Tunable area", "What changes when adjusted", "Classroom use"],
        [
            ["Demand environment", "Total demand, segment shares, price pressure, quality sensitivity.", "Create boom, recession, premium, or price-war rounds."],
            ["Technology pressure", "Market generation and segment adoption of newer tech.", "Reward innovation timing or punish stale portfolios."],
            ["Supply/cost pressure", "Material index, supply risk, supplier cost/defect assumptions.", "Teach cost shocks, sourcing tradeoffs, and quality risk."],
            ["Finance pressure", "Starting cash, interest rate, working-capital burden, liquidity flags.", "Show why profit and survival are not identical."],
            ["Balanced score", "Relative weight on profit, service, forecast, cash, and innovation.", "Prevent one extreme strategy from dominating every lesson."],
        ],
        [1.75, 2.55, 2.1],
        LIGHT_ORANGE,
    )
    add_table(
        doc,
        ["Where it appears", "Main outputs", "What to interpret"],
        [
            ["Team Decisions preview", "Capacity, feasible output, margin, cash, warnings.", "Whether the submitted plan is internally coherent before the round."],
            ["Round engine", "Demand, sales, defects, backlog, inventory, costs, cash.", "How market competition and operational constraints convert decisions into outcomes."],
            ["Results Dashboard", "Rankings, product detail, forecast accuracy, liquidity, lifecycle.", "Which tradeoffs explain performance differences across teams."],
            ["Finance Detail", "Cash bridge, cost buckets, debt, interest, product contribution.", "Why accounting profit, available cash, and survival are different."],
            ["My Team Tracker", "Round history of decisions and results.", "How a team's own pattern evolved over time."],
        ],
        [1.55, 2.7, 2.15],
        LIGHT_BLUE,
    )
    add_table(
        doc,
        ["Formula family", "Main logic", "Why students care"],
        [
            ["Attractiveness", "Products split demand by relative weighted scores.", "Competitors matter; absolute quality is not enough."],
            ["Constraint caps", "Capacity and raw materials limit production.", "A forecast without resources becomes unmet demand."],
            ["Cash bridge", "Revenue and borrowing add cash; costs, NPD, and interest consume cash.", "A profitable plan can still run out of liquidity."],
            ["State carryover", "Inventory, backlog, products, projects, cash, debt, reputation, and capacity persist.", "Each round starts from prior consequences."],
        ],
        [1.55, 2.7, 2.15],
        LIGHT_GREEN,
    )
    add_callout(
        doc,
        "Research note",
        "The coefficients are not claimed as industry estimates. They are transparent teaching parameters that can be calibrated to create different learning environments.",
        LIGHT_PURPLE,
        PURPLE,
    )

    out = DOC_DIR / "Kiki_Simulator_Model_and_Formula_Guide_v2.docx"
    doc.save(out)
    return out


def build_student_guide() -> Path:
    doc = Document()
    style_doc(doc, landscape=True)
    add_title_block(
        doc,
        "Student Guidebook",
        "Kiki Pickleball Business Simulation",
        "Audience: team leaders and student teams | Use this before submitting each round.",
        BLUE,
    )
    add_callout(
        doc,
        "Big idea",
        "You are not only trying to sell paddles. You are managing a business: forecast demand, make products, control quality, manage cash, invest in future products, and learn from each round.",
        LIGHT_BLUE,
        BLUE,
    )

    add_section(doc, "1. What You Control Each Round")
    add_table(
        doc,
        ["Decision area", "Your inputs", "What to think about"],
        [
            ["Market response", "Forecast units, selling price, target segment.", "What do customers want this round?"],
            ["Operations", "Planned production, overtime, capacity expansion, QC budget.", "Can we make enough good units without overspending?"],
            ["Supply", "Raw material order, supplier mix, expedited share.", "Do we want low cost, reliability, or speed?"],
            ["Portfolio", "Product A/B/C active status, price, production, QC, inventory, retirement.", "Which products should carry our business?"],
            ["Pipeline", "NPD investment, testing intensity, launch now when ready.", "Should we spend cash today for future advantage?"],
            ["Finance", "Planned borrowing, cash, debt, working capital.", "Can our strategy survive financially?"],
        ],
        [1.45, 3.0, 4.6],
        LIGHT_GREEN,
    )
    add_callout(
        doc,
        "Decision rhythm",
        "Forecast demand first, then align production, materials, capacity, quality, inventory, pipeline investment, and cash. The best plans fit together.",
        LIGHT_ORANGE,
        ORANGE,
    )
    add_table(
        doc,
        ["Before saving", "Ask your team"],
        [
            ["Demand", "Are our forecasts realistic for our price, segment, product age, and technology?"],
            ["Operations", "Can capacity and raw materials support the production plan?"],
            ["Finance", "Can cash survive inventory, expansion, borrowing, NPD, and interest?"],
            ["Future", "Are we replacing old products before the market moves away from them?"],
        ],
        [1.35, 7.75],
        LIGHT_BLUE,
    )
    add_table(
        doc,
        ["Common mistake", "What usually happens"],
        [
            ["Forecast is zero or unrealistic", "The plan cannot be evaluated well and forecast accuracy suffers."],
            ["Production ignores capacity/materials", "The engine caps output and demand may become backlog or lost sales."],
            ["Price below variable cost", "Sales volume can increase while profit gets worse."],
            ["Heavy expansion/NPD without cash", "Debt and interest rise; liquidity stress can hurt future flexibility."],
            ["Old products never retired", "Declining lifecycle and outdated technology can drag the portfolio."],
        ],
        [2.35, 6.75],
        LIGHT_RED,
    )

    figure_data = [
        (
            "Figure_1_The_Round_Loop.png",
            "The round loop: read the market, submit decisions, see results, then adjust.",
            "Your decisions carry forward. Inventory, backlog, cash, debt, reputation, product age, and pipeline readiness all affect the next round.",
        ),
        (
            "Figure_2_What_Students_Control.png",
            "Control map: the six areas of student decision-making.",
            "Do not optimize one box alone. A good price is useless if you cannot produce, a big launch is dangerous if cash breaks, and high production can backfire if demand is weak.",
        ),
        (
            "Figure_3_Portfolio_and_Pipeline.png",
            "Portfolio and pipeline: active products earn today; projects prepare tomorrow.",
            "Active products compete now. Development projects need funding, testing, timing, and a launch decision before they become sellable products.",
        ),
        (
            "Figure_4_Demand_Allocation_Infographic.png",
            "Demand allocation: products win demand through relative attractiveness.",
            "Your product earns demand based on price fit, quality, segment fit, service, technology, lifecycle, and reputation compared with competitors.",
        ),
        (
            "Figure_5_Operations_and_Cost_Engine.png",
            "Operations engine: capacity, materials, defects, and sales create cost and profit.",
            "Demand is not enough. The simulator caps production by capacity/materials, removes defective units, and then sells only what is available.",
        ),
        (
            "Figure_6_Forecasting_and_SOP.png",
            "Forecasting and S&OP: forecast does not create demand, but it disciplines the plan.",
            "A forecast is your best estimate of demand. If production, materials, capacity, and cash do not match the forecast, your plan may create stockouts, excess inventory, or debt.",
        ),
        (
            "Figure_7_Cash_and_Debt_Pressure.png",
            "Cash and debt: profit and cash are not the same thing.",
            "Borrowing increases cash today but also creates debt and interest. If cash goes negative, the simulator automatically borrows to keep the team alive.",
        ),
        (
            "Figure_8_Strategy_Archetypes.png",
            "Strategy archetypes: there is no permanent winner.",
            "A team can win in one environment and lose in another. The best strategy depends on the market and the scoring method.",
        ),
    ]

    for filename, caption, lesson in figure_data:
        add_figure(doc, FIG_DIR / filename, caption, width=9.55)
        add_callout(doc, "What to remember", lesson, LIGHT_TEAL, TEAL)

    add_section(doc, "2. Decision Page Walkthrough")
    add_numbered(
        doc,
        [
            "Read the Public Market Report first. Total demand, segment shares, material cost, supply risk, quality sensitivity, and technology pressure tell you what kind of round you are entering.",
            "Review current team state. Installed capacity, inventory, backlog, reputation, cash, debt, and tech position are your starting point.",
            "Set firm-level operations and finance. Decide overtime, expansion, raw materials, supplier mix, expedited share, max backorders, and planned borrowing.",
            "Set product-level decisions for Product A/B/C. For each active product, choose price, forecast, production, QC budget, target inventory, and retirement if needed.",
            "Manage the development pipeline. Once a project begins, fixed project settings lock; investment and testing remain adjustable.",
            "Read the Planning Analytics Preview. Fix warnings before saving, especially supplier mix, low raw materials, negative margin, high cash shortfall, and missing forecasts.",
            "Click Save Portfolio and Pipeline Decision before the deadline.",
        ],
    )

    add_section(doc, "3. Variable Meanings")
    add_table(
        doc,
        ["Variable", "What it means", "Common student mistake"],
        [
            ["Forecast units", "Your estimate of demand for each active product.", "Thinking forecast creates demand. It does not."],
            ["Planned production units", "How many units you try to make.", "Producing far more than forecast and tying up cash."],
            ["QC budget per unit", "Spending to reduce defects.", "Setting it too low in picky/high-quality markets."],
            ["Target FG inventory", "Desired finished-goods buffer.", "Too much inventory creates working-capital pressure."],
            ["Supplier mix", "Offshore, balanced, and premium shares; must sum to 100%.", "Forgetting that cheaper supply may increase risk/defects."],
            ["Planned borrowing", "Money borrowed intentionally.", "Using debt as free money; interest and stress follow."],
            ["NPD investment", "Money spent on future product development.", "Expecting instant launch without readiness/timing gates."],
            ["Testing intensity", "Managerial effort to improve readiness and launch stability.", "Investing money but ignoring testing."],
        ],
        [1.65, 4.0, 3.45],
        LIGHT_ORANGE,
    )

    add_section(doc, "4. How Profit and Cost Are Decided")
    add_table(
        doc,
        ["Component", "In plain English"],
        [
            ["Revenue", "Units sold x selling price."],
            ["Procurement cost", "Raw materials purchased, adjusted by supplier mix and material cost index."],
            ["Production cost", "Cost of making planned/actual production units."],
            ["QC cost", "Quality budget per unit times production volume."],
            ["Holding cost", "Cost of leftover finished goods inventory."],
            ["Warranty cost", "Cost created by defects on units sold."],
            ["Backlog cost", "Penalty for unmet demand carried forward."],
            ["Expansion cost", "Capacity investment that mostly helps future rounds."],
            ["NPD cost", "Pipeline investment for future products."],
            ["Interest", "Cost of short-term debt."],
            ["Profit", "Revenue minus all cost components."],
        ],
        [2.0, 7.1],
        LIGHT_RED,
    )
    add_callout(
        doc,
        "Final reminder",
        "A strong submission is not the highest number in every box. It is a coherent plan where forecast, production, sourcing, portfolio, pipeline, and cash all fit the same strategy.",
        LIGHT_TEAL,
        TEAL,
    )

    out = DOC_DIR / "Kiki_Student_Guidebook_with_Figures_v1.docx"
    doc.save(out)
    return out


def main() -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [
        build_instructor_manual(),
        build_model_guide(),
        build_student_guide(),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
