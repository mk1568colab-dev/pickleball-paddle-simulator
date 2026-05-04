"""Build detailed Kiki classroom documentation pack.

This script generates the longer instructor, model/formula, and student guides.
It intentionally keeps the docs separate from the older compact Kiki pack so
the instructor can choose either the quick version or the detailed version.
"""

from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "documentation"
FIG_DIR = DOC_DIR / "figures" / "game_big_picture"
KIKI_IMAGE = ROOT / "assets" / "kiki_mascot.png"
SCREEN_DIR = Path(r"C:/Users/mk156/Pictures/Screenshots")
EXPERIMENT_SCRIPT = ROOT / "scripts" / "run_20_environment_simulations.py"
STRATEGY_RUNNER_SCRIPT = ROOT / "scripts" / "run_strategy_simulation.py"
REQUIREMENTS_FILE = ROOT / "requirements.txt"
SIMULATION_OUTPUTS_DIR = ROOT / "simulation_outputs"

BLUE = "2563EB"
GREEN = "16A34A"
ORANGE = "D97706"
PURPLE = "7C3AED"
RED = "DC2626"
TEAL = "0D9488"
INK = "172033"
MUTED = "59677B"
LIGHT_BLUE = "DBEAFE"
LIGHT_GREEN = "DCFCE7"
LIGHT_ORANGE = "FEF3C7"
LIGHT_PURPLE = "EDE9FE"
LIGHT_RED = "FEE2E2"
LIGHT_TEAL = "CCFBF1"
LIGHT_GRAY = "F8FAFC"


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _margins(cell, margin: int = 115) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def _widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def _compact(paragraph, before: int = 0, after: int = 4, line: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def _style_doc(doc: Document, *, landscape: bool = False) -> None:
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
    else:
        section.left_margin = Inches(0.68)
        section.right_margin = Inches(0.68)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)

    styles = doc.styles
    for style_name in ("Normal", "List Bullet", "List Number"):
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(9.6)
        styles[style_name].font.color.rgb = RGBColor.from_string(INK)
    for style_name, size in (
        ("Title", 22),
        ("Heading 1", 15.5),
        ("Heading 2", 12.2),
        ("Heading 3", 10.7),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
    styles["Caption"].font.name = "Aptos"
    styles["Caption"].font.size = Pt(8.4)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor.from_string(MUTED)


def _usable_width(doc: Document) -> float:
    section = doc.sections[0]
    return section.page_width.inches - section.left_margin.inches - section.right_margin.inches


def _title_block(
    doc: Document,
    title: str,
    subtitle: str,
    audience: str,
    *,
    accent: str,
    mascot: bool = True,
) -> None:
    has_mascot = mascot and KIKI_IMAGE.exists()
    table = doc.add_table(rows=1, cols=2 if has_mascot else 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    usable = _usable_width(doc)
    _widths(table, [usable - 1.65, 1.55] if has_mascot else [usable])
    for cell in table.rows[0].cells:
        _shade(cell, LIGHT_GRAY)
        _margins(cell, 150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = table.rows[0].cells[0]
    paragraph = left.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor.from_string(accent)
    _compact(paragraph, after=5)

    paragraph = left.add_paragraph(subtitle)
    paragraph.runs[0].font.size = Pt(10.5)
    paragraph.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    _compact(paragraph, after=5)

    paragraph = left.add_paragraph(audience)
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(9.2)
    _compact(paragraph, after=0)

    if has_mascot:
        paragraph = table.rows[0].cells[1].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(KIKI_IMAGE), width=Inches(1.25))
        _compact(paragraph, after=0)


def _callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = LIGHT_BLUE,
    accent: str = BLUE,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _widths(table, [_usable_width(doc)])
    cell = table.rows[0].cells[0]
    _shade(cell, fill)
    _margins(cell, 130)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(accent)
    paragraph.add_run(text)
    _compact(paragraph, after=0)


def _h1(doc: Document, title: str, intro: str | None = None) -> None:
    heading = doc.add_heading(title, level=1)
    _compact(heading, before=7, after=4)
    if intro:
        paragraph = doc.add_paragraph(intro)
        paragraph.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        _compact(paragraph, after=5)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        _compact(paragraph, after=1)


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)
        _compact(paragraph, after=1)


def _table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float],
    *,
    fill: str = LIGHT_BLUE,
    font_size: float = 8.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _widths(table, col_widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade(cell, fill)
        _margins(cell, 90)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(INK)
        _compact(paragraph, after=0, line=1.0)

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            _margins(cell, 90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.add_run(str(value))
            paragraph.runs[0].font.size = Pt(font_size - 0.25)
            _compact(paragraph, after=0, line=1.0)


def _figure(doc: Document, path: Path, caption: str, *, width: float) -> bool:
    if not path.exists():
        _callout(
            doc,
            "Screenshot needed",
            f"Missing image: {path}",
            fill=LIGHT_ORANGE,
            accent=ORANGE,
        )
        return False

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    _compact(paragraph, after=1)

    caption_paragraph = doc.add_paragraph(caption, style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _compact(caption_paragraph, after=3)
    return True


def _figure_explanation(
    doc: Document,
    title: str,
    purpose: str,
    how_to_read: str,
    use: str,
    student_note: str | None = None,
    *,
    fill: str = LIGHT_TEAL,
) -> None:
    rows = [
        ["Purpose", purpose],
        ["How to read it", how_to_read],
        ["Teaching/use", use],
    ]
    if student_note:
        rows.append(["Student note", student_note])
    _table(
        doc,
        [f"{title}", "Explanation"],
        rows,
        [1.35, _usable_width(doc) - 1.35],
        fill=fill,
        font_size=8.0,
    )


def build_instructor_manual() -> Path:
    doc = Document()
    _style_doc(doc)
    _title_block(
        doc,
        "Instructor Operating Manual",
        "Kiki Pickleball Business Simulation",
        "Detailed classroom operations guide for instructors and TAs | Updated for Kiki branding, finance detail, team tracker, product pipeline, and factory reset.",
        accent=PURPLE,
    )
    _callout(
        doc,
        "What this file is",
        "A practical runbook for setting up accounts, controlling rounds, explaining the UI, validating submissions, running the engine, and debriefing results. It is intentionally more detailed than the quick-start checklist.",
        fill=LIGHT_PURPLE,
        accent=PURPLE,
    )
    _table(
        doc,
        ["Manual section", "What it helps you do"],
        [
            ["Setup", "Deploy/open the app, create accounts, confirm roles, test the database, and prepare the first market report."],
            ["Round control", "Open submissions, monitor teams, close submissions, run the round, and advance to the next market report."],
            ["Teaching support", "Use dashboards, finance detail, team tracker, formula guide, and big-picture figures to connect outcomes to OM/SCM concepts."],
            ["Troubleshooting", "Know what to do when students miss forecasts, supplier mix is invalid, teams cannot log in, or a reset is needed."],
        ],
        [1.6, 4.95],
        fill=LIGHT_BLUE,
    )
    _callout(
        doc,
        "Important",
        "Students should use the hosted Render URL during class, not http://localhost:8501. Localhost only works on the computer running Streamlit.",
        fill=LIGHT_RED,
        accent=RED,
    )

    _h1(
        doc,
        "1. Simulator Purpose and Teaching Goals",
        "Kiki is a classroom business simulation, not just a scoreboard game. Teams run paddle companies and learn how decisions interact over repeated rounds.",
    )
    _table(
        doc,
        ["Concept", "What students practice", "Typical discussion question"],
        [
            ["Demand forecasting", "Estimate product demand before production and compare forecast vs actual.", "Did the team forecast demand or simply guess a production number?"],
            ["Market segmentation", "Compete in premium, mid, and beginner segments with different price/quality/tech expectations.", "Which segment fit the team strategy best this round?"],
            ["Product portfolio", "Manage up to three active product slots with different lifecycle, price, QC, and production choices.", "Did the portfolio contain the right products for the market?"],
            ["Lifecycle management", "Products move from launch to growth to maturity to decline.", "Which product should be invested in, maintained, retired, or replaced?"],
            ["Capacity planning", "Balance installed capacity, overtime, and expansion.", "Did growth create capability or just cost and debt?"],
            ["Supply chain", "Choose raw-material quantity, supplier mix, and expedited share.", "Did cheaper supply create hidden quality or service risk?"],
            ["Quality control", "Spend QC budget to lower defects with diminishing returns.", "Was quality worth the cost in this market?"],
            ["Inventory/backlog", "Carry finished goods, raw materials, and unmet demand across rounds.", "Did inventory protect service or tie up too much cash?"],
            ["Cash and debt", "Manage revenue, costs, planned borrowing, automatic borrowing, and interest.", "Was the team profitable but still financially fragile?"],
            ["NPD and technology", "Invest in future products, wait for readiness, and launch/replace when ready.", "Did the team innovate early enough without starving current operations?"],
        ],
        [1.4, 2.65, 2.5],
        fill=LIGHT_GREEN,
        font_size=7.7,
    )

    _h1(doc, "2. User Roles and Login Workflow")
    _table(
        doc,
        ["Role", "Access", "What they can do", "What they cannot do"],
        [
            ["admin", "All pages.", "Create users, set market report, view all teams, open/close submissions, run rounds, reset game data, inspect finance/results.", "Should not be shared with students."],
            ["team_leader", "Student-safe pages only.", "Submit only their assigned team decisions, view public report, view own tracker/finance/results and safe rankings.", "Cannot run rounds, reset data, manage users, or edit other teams."],
        ],
        [1.0, 1.55, 2.35, 1.65],
        fill=LIGHT_BLUE,
    )
    _bullets(
        doc,
        [
            "First-run setup appears only when no active admin exists. After an admin is created, normal login is used.",
            "Passwords are hashed. The admin can reset a password but cannot view an old password later.",
            "Each team should use its own team_leader account so decisions and results remain separated.",
        ],
    )
    _figure(doc, SCREEN_DIR / "first page.png", "Figure 1. Home page after login. The sidebar shows navigation, logged-in user, role, and logout.", width=5.9)
    _figure_explanation(
        doc,
        "Home page",
        "Orient the instructor and confirm that the correct account is logged in.",
        "The left sidebar is the control rail. Home explains app capabilities and points instructors to the pages used each round.",
        "At the start of class, project this page briefly to show students that the app is a central hosted classroom tool.",
    )

    _h1(doc, "3. Local Testing vs Hosted Classroom Deployment")
    _table(
        doc,
        ["Mode", "Typical URL", "Who should use it", "Notes"],
        [
            ["Local testing", "http://localhost:8501", "Instructor/developer on the same computer.", "Good for testing. Students on other devices usually cannot use this URL."],
            ["Hosted classroom", "https://your-render-service.onrender.com", "Instructor and students from browsers.", "Use this for class. Students can join from school Wi-Fi, conference Wi-Fi, home, or cellular data."],
        ],
        [1.25, 2.1, 1.8, 1.4],
        fill=LIGHT_ORANGE,
    )
    _bullets(
        doc,
        [
            "Render deployment uses one central Streamlit service and a persistent SQLite database path, usually SIMULATOR_DB_PATH=/var/data/simulator.db.",
            "Use one hosted app instance for the class. SQLite is appropriate for a small single-instance classroom deployment, not multiple replicas writing to separate disks.",
            "After pushing code to GitHub, Render redeploys from the repo/blueprint. The live class URL is the Render service URL, not the GitHub repo URL.",
        ],
    )

    _h1(doc, "4. Instructor Pre-Class Setup Checklist")
    _table(
        doc,
        ["Done", "Task", "Why it matters"],
        [
            ["[ ]", "Open the hosted URL and log in as admin.", "Confirms the deployment is alive."],
            ["[ ]", "Verify database path uses persistent Render disk.", "Avoids losing data after restart."],
            ["[ ]", "Create or verify team_leader accounts.", "Every team needs its own login and team assignment."],
            ["[ ]", "Test one student account in a private/incognito browser.", "Confirms student restrictions and login work."],
            ["[ ]", "Set Public Market Report for Round 1.", "Defines the environment before decisions."],
            ["[ ]", "Open submissions.", "Allows teams to save decisions."],
            ["[ ]", "Run one private test round before live use.", "Catches deployment/data/account issues early."],
        ],
        [0.55, 2.4, 3.6],
        fill=LIGHT_GREEN,
    )

    _h1(doc, "5. Page-by-Page Instructor Guide")
    _table(
        doc,
        ["Page", "Admin use", "Student/team_leader view"],
        [
            ["Home", "Confirm version, role, database, and page navigation.", "Orientation, assigned team, student workflow."],
            ["Public Market Report", "Edit market, demand, tech, quality, supply, and event text.", "Read-only market conditions."],
            ["Team Decisions", "Can inspect/enter any team if needed.", "Only their assigned team; submit firm, product, forecast, and pipeline decisions."],
            ["Instructor Panel", "Admin-only validation, submissions, run round, factory reset.", "No access."],
            ["Admin User Management", "Admin-only account creation, reset, deactivate, bulk import.", "No access."],
            ["Results Dashboard", "All teams, rankings, product results, debrief tabs, downloads.", "Own team details plus safe public rankings."],
            ["Finance Detail", "Detailed finance by team/round; cash bridge and cost buckets.", "Own finance only."],
            ["My Team Tracker", "Admin can inspect a selected team history.", "Own decisions/results history."],
            ["Game Big Picture Guide", "Use as teaching intro/debrief visual guide.", "Student-friendly system explanation."],
            ["Model Formula Guide", "Explain deterministic formulas and coefficients.", "Student reference for how outcomes are calculated."],
            ["My Account", "Change own password.", "Change own password."],
        ],
        [1.3, 2.65, 2.6],
        fill=LIGHT_BLUE,
        font_size=7.65,
    )

    _h1(doc, "6. Public Market Report Setup", "Set or apply the market scenario before teams submit decisions. This page defines the world all teams compete in for the round.")
    _figure(doc, SCREEN_DIR / "public market report.png", "Figure 2. Public Market Report page: demand, segment shares, supply/cost pressure, quality sensitivity, technology pressure, and market event.", width=6.15)
    _figure_explanation(
        doc,
        "Public Market Report",
        "Create the shared environment for the round.",
        "Demand and segment shares create market size. Material cost index and supply risk affect operations. Technology and adoption settings change how valuable newer products are.",
        "Use the event text to tell a story: recession, boom, technology shift, price war, quality scandal, or beginner growth.",
        fill=LIGHT_PURPLE,
    )
    _table(
        doc,
        ["Control", "Meaning", "Instructor guidance"],
        [
            ["Round Number", "The round being configured.", "After Run Round, the app automatically prepares the next round."],
            ["Total Demand", "Total potential market units before competition allocation.", "Higher values create growth/boom conditions; lower values create recession."],
            ["Premium/Mid/Beginner Share", "How demand is split across customer segments.", "Shares should sum close to 1.0. Changing shares changes which strategies are advantaged."],
            ["Material Cost Index", "Multiplier on material costs.", "Use high values for inflation or supply cost shock."],
            ["Supply Risk", "Low/Moderate/High risk pressure on defects/reliability.", "High supply risk rewards robust sourcing, QC, and conservative planning."],
            ["Quality Sensitivity", "How strongly customers/reputation react to quality.", "Higher values reward QC and punish defects."],
            ["Current Market Generation", "Technology generation customers expect.", "Higher generation makes old tech less attractive, especially premium/mid."],
            ["Technology Shift Rate", "How quickly the market moves toward newer products.", "Use high values to reward NPD and innovation timing."],
            ["Premium/Mid Tech Adoption", "How much premium/mid customers care about newer tech.", "Premium usually cares more than beginner."],
            ["Beginner Price Pressure", "How strongly beginner customers react to price.", "High values reward low-cost/value strategies."],
            ["Market Event", "Narrative text shown to all teams.", "Use it to explain the scenario and guide student interpretation."],
            ["Save Market Report", "Stores the environment for the round.", "Save before teams submit. If you change it mid-round, tell the class."],
        ],
        [1.45, 2.5, 2.6],
        fill=LIGHT_PURPLE,
        font_size=7.35,
    )

    _h1(doc, "7. Team Decisions Page Guide", "Students submit a coherent plan, not isolated high/low choices. The instructor can use this page to explain what a good decision sequence looks like.")
    _numbered(
        doc,
        [
            "Read the Public Market Report.",
            "Review current team state.",
            "Set firm-level operations and finance.",
            "Review and update active product slots.",
            "Enter product-level forecasts, prices, production, QC, and inventory targets.",
            "Invest in or manage development projects.",
            "Review Planning Analytics Preview and warnings.",
            "Click Save Portfolio and Pipeline Decision.",
        ],
    )
    _figure(doc, SCREEN_DIR / "team decision (1).png", "Figure 3. Team Decisions upper section: identity, current state, firm-level operations and finance.", width=6.15)
    _figure_explanation(
        doc,
        "Team Decisions upper section",
        "Shows where students start and what resources they carry into the round.",
        "Installed capacity, FG/RM inventory, backlog, reputation, cash, debt, and tech position are state variables from prior rounds. Firm controls affect the whole company, not one product.",
        "Use this section to teach capacity, materials, borrowing, and supplier mix as shared constraints across products.",
    )
    _figure(doc, SCREEN_DIR / "team decision (2).png", "Figure 4. Active Product Portfolio and Forecasts: product A/B/C decisions.", width=6.15)
    _figure_explanation(
        doc,
        "Active portfolio",
        "Each active product has its own forecast, price, production, QC, target inventory, lifecycle, and tech generation.",
        "Products compete at the product level, then results roll up to the firm. Retire After Round flags an aging product for retirement after the round logic runs.",
        "Ask teams why each product exists: premium margin, mid-market volume, beginner volume, or replacement bridge.",
    )
    _figure(doc, SCREEN_DIR / "team decision (3).png", "Figure 5. Development Pipeline: two project slots for future products.", width=6.15)
    _figure_explanation(
        doc,
        "Development pipeline",
        "Future products require investment, readiness, and launch timing before they are sellable.",
        "Fixed project settings lock once a project begins: name, segment, target technology generation, intended slot, and launch timing. Round controls remain adjustable: investment, testing intensity, launch now if ready, or cancel.",
        "Emphasize that more money can speed funding progress, but readiness also needs testing and timing gates.",
        fill=LIGHT_PURPLE,
    )
    _figure(doc, SCREEN_DIR / "team decision (4).png", "Figure 6. Planning Analytics Preview: pre-submit diagnostics.", width=6.15)
    _figure_explanation(
        doc,
        "Planning preview",
        "The page translates inputs into planning signals before save.",
        "Forecast-production gap, utilization, raw material sufficiency, weighted cost, defect estimate, projected margin, cash before borrowing, and warnings tell students whether the plan fits together.",
        "Tell students: if the preview is warning you, fix the plan before saving.",
        fill=LIGHT_ORANGE,
    )
    _table(
        doc,
        ["Student input", "Why it exists", "Instructor teaching point"],
        [
            ["Forecast Units", "Demand estimate for each active product.", "Forecast does not create demand; it measures planning discipline."],
            ["Planned Production Units", "Attempted production by product.", "Production is capped by capacity and raw materials."],
            ["Selling Price / Unit", "Revenue per unit and price competitiveness.", "Price affects both margin and demand attractiveness."],
            ["QC Budget / Unit", "Quality investment that lowers defects with diminishing returns.", "Quality is valuable but not free."],
            ["Target FG Inventory", "Desired finished-goods buffer.", "Inventory protects service but ties up working capital."],
            ["Overtime Capacity", "Short-term capacity above installed base.", "Can help service but creates cost/stress."],
            ["Capacity Expansion", "Future installed capacity investment.", "Growth takes cash before it pays off."],
            ["Raw Material Order", "New material supply for production.", "Low materials can cap output; too much ties cash."],
            ["Supplier Mix", "Offshore/balanced/premium shares.", "Cost, lead time, risk, and defect pressure trade off."],
            ["Planned Borrowing", "Intentional short-term cash injection.", "Debt helps liquidity now but creates interest."],
        ],
        [1.45, 2.5, 2.6],
        fill=LIGHT_ORANGE,
        font_size=7.35,
    )

    _h1(doc, "8. Development Pipeline Guide")
    _table(
        doc,
        ["Project concept", "Rule in Kiki", "How to explain it"],
        [
            ["Fixed project settings", "After investment/testing begins, project name, target segment, target tech generation, intended slot, and planned launch timing lock.", "This models product concept freeze. Teams can still execute the project, but not rewrite the original concept every round."],
            ["Controllable project decisions", "Investment this round, testing intensity, launch now if ready, and cancel project remain adjustable.", "These are managerial control decisions during development."],
            ["Readiness", "Readiness increases with cumulative investment and testing; higher tech requires more funding and usually more effort.", "Money alone may not instantly finish a project. Testing and timing gates matter."],
            ["Launch gate", "Project must be sufficiently funded, readiness-ready, and eligible by timing before launch.", "Students should plan ahead; waiting until the market shifts may be too late."],
            ["Replacement", "A launched project can fill or replace an intended product slot.", "Replacement refreshes the portfolio but can cannibalize older products."],
            ["Retirement", "Retired products stop competing for demand after retirement is processed.", "Retire products that are old, low-tech, low-margin, or no longer strategic."],
        ],
        [1.6, 2.65, 2.3],
        fill=LIGHT_PURPLE,
        font_size=7.6,
    )
    _callout(
        doc,
        "Instructor tip",
        "If students say readiness is stuck, ask whether funding is enough for the tech generation, whether testing intensity is meaningful, whether the timing gate has arrived, and whether they checked Launch Now If Ready.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    _h1(doc, "9. Instructor Panel and Round Execution")
    _figure(doc, SCREEN_DIR / "instructor panel.png", "Figure 7. Instructor Panel top: validation metrics, submission controls, and round controls.", width=6.15)
    _figure_explanation(
        doc,
        "Instructor Panel top",
        "Run the classroom round safely.",
        "The top metrics warn about invalid supplier mix, infeasible plans, missing forecasts, forecast mismatches, cash shortfalls, pipeline projects, and launch-ready projects.",
        "Use this as the pre-flight checklist before clicking Run Round.",
        fill=LIGHT_RED,
    )
    _figure(doc, SCREEN_DIR / "instructor panel _2.png", "Figure 8. Instructor Panel detail: validation detail, saved decisions, projects, and portfolio snapshot.", width=6.15)
    _figure_explanation(
        doc,
        "Instructor Panel detail",
        "Audit exactly what has been submitted before running the engine.",
        "Saved Firm-Level Decisions confirms shared operations choices. Saved Product-Level Decisions shows product/forecast inputs. Saved Development Projects shows pipeline state. Portfolio Snapshot shows current products.",
        "When students dispute a result, start here: did they actually save what they thought they saved?",
        fill=LIGHT_RED,
    )
    _table(
        doc,
        ["Control/metric", "Meaning", "Action"],
        [
            ["Teams Submitted", "Number of teams with saved firm-level decisions.", "Check against expected active teams."],
            ["Invalid Supplier Mix", "Supplier shares do not sum to 100%.", "Ask team to fix; engine can normalize, but it is a planning error."],
            ["Infeasible Plans", "Production exceeds shared capacity/material feasibility.", "Warn team that output may be capped."],
            ["Missing Forecasts", "Active products have zero/missing forecasts.", "Require teams to enter forecasts before run if forecasting is graded."],
            ["Likely Cash Shortfalls", "Submission likely requires borrowing.", "Use as finance teaching moment; not always an error."],
            ["Open Team Submissions", "Lets teams save decisions.", "Use at start of decision period or for late fixes."],
            ["Close Team Submissions", "Locks normal submission window.", "Use before final pre-run check."],
            ["Run Round", "Executes the OM/SCM/portfolio/finance engine.", "Click after submissions are acceptable; app opens next round automatically."],
            ["Factory Reset Game", "Erases game progress but keeps accounts.", "Use only for testing or intentional restart."],
        ],
        [1.55, 2.55, 2.45],
        fill=LIGHT_RED,
        font_size=7.35,
    )
    _callout(
        doc,
        "Warning",
        "Factory Reset erases game progress, market report history, decisions, product portfolios, pipeline progress, results, and team states. It keeps user accounts. Do not use during a live class unless you intentionally want to restart the game.",
        fill=LIGHT_RED,
        accent=RED,
    )

    _h1(doc, "10. Admin User Management")
    _figure(doc, SCREEN_DIR / "admin management.png", "Figure 9. Admin User Management: account table, filters, create/edit/bulk import tabs.", width=6.15)
    _figure_explanation(
        doc,
        "Admin User Management",
        "Create and maintain instructor and team leader accounts.",
        "The table lists usernames, role, team name, and active status but never password hashes or old passwords. Use Create User for one account, Edit Existing User to reset/deactivate, and Bulk Import for many teams.",
        "One account per team is cleaner than everyone sharing admin or one shared class account.",
    )
    _table(
        doc,
        ["Task", "How to do it", "Important note"],
        [
            ["Create admin", "Create User tab, role=admin, blank team name.", "Only trusted instructors/TAs should be admin."],
            ["Create team leader", "Create User tab, role=team_leader, assign team name.", "Team name is required and drives data filtering."],
            ["Reset password", "Edit Existing User tab, enter a new temporary password.", "Password is shown only at reset/create time; it is not retrievable later."],
            ["Deactivate user", "Edit Existing User, uncheck Active.", "Deactivated users cannot log in."],
            ["Bulk import teams", "Upload CSV with username,password,team_name,is_active.", "Good for creating 6-12 teams quickly."],
            ["Remove team", "Deactivate the account or use reset/cleanup workflow depending on classroom need.", "Keep account history if you need audit trail. Factory reset keeps accounts but clears game data."],
        ],
        [1.4, 2.85, 2.3],
        fill=LIGHT_BLUE,
        font_size=7.55,
    )

    _h1(doc, "11. Results Dashboard and Teaching Debrief")
    _figure(doc, SCREEN_DIR / "results dashboard.png", "Figure 10. Results Dashboard: market report, tabs, portfolio snapshot, lifecycle distribution, and downloads.", width=6.15)
    _figure_explanation(
        doc,
        "Results Dashboard",
        "Debrief results after the round.",
        "The top repeats the market report. Tabs separate team summary, teaching debrief, forecast accuracy, product detail, liquidity/debt, portfolio snapshot, development pipeline, and launch/retirement log.",
        "Use it to explain why winners won, not just who won.",
    )
    _table(
        doc,
        ["Dashboard tab", "What it shows", "Debrief question"],
        [
            ["Team Summary", "Team-level revenue, cost, profit, service, cash, debt, and rankings.", "Was the winner profitable, balanced, or just lucky in one dimension?"],
            ["Teaching Debrief", "Plain-English diagnostic patterns and class discussion prompts.", "What tradeoff does this round best illustrate?"],
            ["Forecast Accuracy", "Forecast vs actual by product/team, error, WAPE/bias.", "Which teams planned based on demand, and which teams guessed?"],
            ["Product Detail", "Product-level demand, sales, defects, inventory, backlog, profit contribution.", "Which product carried the company? Which product should be retired?"],
            ["Liquidity / Debt", "Cash, debt, interest, borrowing, liquidity stress.", "Did growth create cash pressure?"],
            ["Portfolio Snapshot", "Current active/inactive products, lifecycle, tech generation, inventory/backlog.", "Is the portfolio aging or ready for the next market?"],
            ["Development Pipeline", "NPD project status, investment, readiness, launch eligibility.", "Who is investing ahead of the next technology shift?"],
            ["Launch / Retirement Log", "Products launched or retired this round.", "Did replacement timing help or hurt?"],
        ],
        [1.45, 2.7, 2.4],
        fill=LIGHT_TEAL,
        font_size=7.35,
    )

    _h1(doc, "12. Finance Detail, My Team Tracker, Game Big Picture, and Formula Guide")
    _table(
        doc,
        ["Page", "Purpose", "How to use in class"],
        [
            ["Finance Detail", "Shows cash bridge, cost breakdown, product contribution, debt, interest, and liquidity diagnostics.", "Use when students ask why profit and cash differ, or why a high-sales team has debt."],
            ["My Team Tracker", "Shows a team round-by-round history of saved decisions, product decisions, results, forecasts, portfolio, projects, and launch history.", "Use for student reflection: what changed from Round 1 to Round 5?"],
            ["Game Big Picture Guide", "Student-facing visual system guide using the eight Kiki diagrams.", "Use before the first round and again after students feel overwhelmed."],
            ["Model Formula Guide", "Transparent formula reference: demand allocation, attractiveness, cost, forecast, finance, development, and scoring logic.", "Use to explain the model is deterministic and rule-based, not an AI black box."],
        ],
        [1.4, 2.65, 2.5],
        fill=LIGHT_PURPLE,
        font_size=7.55,
    )
    _figure(doc, SCREEN_DIR / "model guide.png", "Figure 11. Model Formula Guide: formula explanation page for deterministic engine logic.", width=5.75)
    _figure_explanation(
        doc,
        "Model Formula Guide",
        "Make the model visible so students can debate assumptions and learn the formulas.",
        "The page shows segment demand, proportional allocation, attractiveness score families, and the logic behind financial and operations outputs.",
        "Use it to shift discussion from gaming the app to reasoning about managerial tradeoffs.",
        fill=LIGHT_BLUE,
    )
    _callout(
        doc,
        "Screenshot status",
        "Current screenshots cover the main v6 UI plus Team Decisions, Instructor Panel, Results Dashboard, Admin Management, Public Market Report, Home, and Formula Guide. New screenshots would be useful only if you want the manual to show the latest Kiki-branded Login page, Finance Detail page, My Team Tracker page, or Game Big Picture page exactly.",
        fill=LIGHT_ORANGE,
        accent=ORANGE,
    )

    _h1(doc, "13. Recommended Classroom Round Routine")
    _table(
        doc,
        ["Phase", "Instructor actions", "Student actions"],
        [
            ["Before class", "Open hosted app, verify DB path, create accounts, set market report, open submissions, test one student login.", "None or review account credentials."],
            ["Decision time", "Ask teams to read market report; monitor Instructor Panel validation; remind teams to save.", "Forecast, set firm/product/pipeline decisions, read preview warnings, save."],
            ["Before run", "Close submissions; check missing forecasts, invalid mix, infeasible plans, cash shortfalls; decide on late fixes.", "Make final fixes if allowed."],
            ["Run/debrief", "Click Run Round; open Results Dashboard, Finance Detail, and Teaching Debrief.", "Compare decisions vs outcomes; explain forecast, service, cash, product, and pipeline results."],
            ["Prepare next round", "Public Market Report is ready for next round; adjust scenario; open submissions.", "Review prior result and adapt strategy."],
        ],
        [1.2, 3.35, 2.0],
        fill=LIGHT_GREEN,
        font_size=7.35,
    )

    _h1(doc, "14. Common Student Mistakes")
    _table(
        doc,
        ["Mistake", "Likely outcome", "Instructor coaching response"],
        [
            ["Using zero forecasts", "Forecast accuracy meaningless or poor; validation warning.", "Forecast first, then plan production."],
            ["Supplier mix not 100%", "Engine normalizes; indicates poor planning.", "Make them explain sourcing strategy."],
            ["Price below variable cost", "High sales can still lose money.", "Ask for unit margin, not just units sold."],
            ["Overproducing demand", "Inventory, holding cost, working capital pressure.", "Ask whether inventory was a buffer or a mistake."],
            ["Ignoring cash while expanding/NPD", "Debt, interest, liquidity warning.", "Separate profit from cash."],
            ["Never retiring products", "Aging tech/lifecycle drag.", "Ask which product should be replaced."],
            ["Launching too late", "Market tech moves before product is ready.", "Plan NPD lead time earlier."],
            ["Focusing only on profit rank", "Misses forecast, service, cash, and innovation learning.", "Use balanced score/debrief prompts."],
        ],
        [1.55, 2.4, 2.6],
        fill=LIGHT_RED,
        font_size=7.35,
    )

    _h1(doc, "15. Troubleshooting Guide")
    _table(
        doc,
        ["Problem", "Likely cause", "Fix"],
        [
            ["Student cannot log in", "Wrong username/password, inactive user, wrong URL.", "Check Admin User Management; reset password; confirm hosted URL."],
            ["Student sees wrong team", "Account assigned to wrong team_name.", "Edit user team assignment."],
            ["Team Decisions not saving", "Submissions closed, validation issue, session stale.", "Open submissions; refresh; check warnings; save again."],
            ["Round will not run", "No team decisions saved.", "At least one team must save firm-level decision."],
            ["Public Market Report still old", "Round not advanced or page not refreshed.", "After Run Round, report advances; refresh page if needed."],
            ["Need complete restart", "Testing/class restart.", "Use Factory Reset Game; accounts remain."],
            ["Render site has database column error", "Old SQLite persistent disk schema vs new code.", "Run migrations/redeploy; if needed reset data or recreate persistent DB after backup."],
            ["Students use localhost", "They copied local testing URL.", "Give Render URL only."],
        ],
        [1.55, 2.2, 2.8],
        fill=LIGHT_ORANGE,
        font_size=7.35,
    )

    _h1(doc, "16. Quick Start Checklist")
    _bullets(
        doc,
        [
            "Before class: hosted URL works, admin login works, teams exist, Round 1 market report saved, submissions open.",
            "During decisions: teams read market report, enter forecasts, align production/material/cash, save final decision.",
            "Before run: close submissions, check validation metrics, decide whether late fixes are allowed.",
            "After run: dashboard debrief, finance detail discussion, team tracker reflection, next round market report adjustment.",
        ],
    )

    out = DOC_DIR / "Kiki_Instructor_Operating_Manual_v4_DETAILED.docx"
    doc.save(out)
    return out


def build_model_guide() -> Path:
    doc = Document()
    _style_doc(doc)
    _title_block(
        doc,
        "Simulator Model and Formula Guide",
        "Kiki Pickleball Business Simulation",
        "Detailed technical and teaching explanation of how the deterministic OM/SCM engine converts decisions into outcomes.",
        accent=BLUE,
    )
    _callout(
        doc,
        "Model philosophy",
        "The simulator is rule-based and transparent. Coefficients are teaching calibration choices, not claims that the real pickleball market follows exact constants.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    _h1(doc, "1. What the Model Represents")
    _table(
        doc,
        ["Layer", "State/decision variables", "Learning purpose"],
        [
            ["Market", "Demand, segment shares, supply risk, material cost index, quality sensitivity, technology pressure.", "Students face environments that reward different strategies."],
            ["Firm operations", "Capacity, overtime, expansion, production, raw materials, supplier mix, expedited share.", "Capacity and materials constrain execution."],
            ["Product portfolio", "Up to three active product slots with price, forecast, production, QC, inventory, lifecycle, tech.", "Products compete separately but share firm constraints."],
            ["NPD pipeline", "Up to two projects with funding, testing, readiness, launch timing, tech generation, target slot.", "Future products require lead time and investment."],
            ["Finance", "Cash, revenue, procurement, production, QC, holding, warranty, backlog, expansion, NPD, interest, debt.", "Profit, cash, and survival are related but different."],
            ["Performance", "Demand, sales, fill rate, forecast accuracy, profit, balanced score, reputation, innovation position.", "Winners depend on the metric and environment."],
        ],
        [1.35, 3.3, 1.9],
        fill=LIGHT_GREEN,
        font_size=7.5,
    )

    _h1(doc, "2. Round Algorithm")
    _numbered(
        doc,
        [
            "Load the current public market report.",
            "Load team decisions, product decisions, forecasts, active product lines, development projects, and persistent team states.",
            "Initialize missing team state and default product portfolio from archetype baselines.",
            "Compute shared firm constraints: installed capacity, overtime, raw materials, supplier cost, supplier lead time, supplier defect pressure, and cash.",
            "Progress NPD projects using investment and testing; determine whether each project is concept, development, pilot, launch_ready, launched, or canceled.",
            "Apply launch, retirement, and replacement choices when gates are satisfied.",
            "Allocate feasible production across active product slots under shared capacity and raw-material constraints.",
            "Compute product-level defect rates and good units produced.",
            "Compute product attractiveness for premium, mid, and beginner market segments.",
            "Allocate segment demand across all active products from all teams in proportion to attractiveness.",
            "Apply conservative intra-team cannibalization when newer/same-group products steal demand from older products.",
            "Realize sales, lost sales/backlog, ending inventory, revenue, cost, profit, and contribution margin.",
            "Compute product and team forecast accuracy.",
            "Update cash, debt, interest, working capital stress, reputation, inventory, backlog, capacity, product age, lifecycle, and pipeline readiness.",
            "Save team-level results, product-level results, forecasts, states, products, and projects to SQLite.",
        ],
    )

    _h1(doc, "3. Demand and Attractiveness")
    _table(
        doc,
        ["Formula family", "Plain formula", "Explanation"],
        [
            ["Segment demand", "premium_demand = total_demand * premium_share; mid_demand = total_demand * mid_share; beginner_demand = total_demand * beginner_share.", "The public market report creates the demand pool."],
            ["Product demand share", "product_segment_demand = segment_demand * product_attractiveness / total_attractiveness_in_segment.", "Demand is allocated by relative attractiveness, not by a fixed sales quota."],
            ["Attractiveness score", "weighted price + quality + segment fit + product fit + reputation + service readiness, then adjusted by lifecycle and technology.", "A product with better fit gets more demand, but competitors matter."],
            ["Price competitiveness", "Actual selling price is compared with segment reference prices.", "No coarse price tier; numeric price matters."],
            ["Technology modifier", "Product tech generation vs current market generation, weighted by segment tech adoption.", "Premium/mid customers usually value newer tech more than beginner customers."],
            ["Lifecycle modifier", "Launch novelty, growth strength, maturity stability, decline penalty.", "Age changes demand potential over time."],
        ],
        [1.35, 3.1, 2.1],
        fill=LIGHT_BLUE,
        font_size=7.35,
    )
    _callout(
        doc,
        "Theory link",
        "This is a weighted multi-attribute attraction model with proportional demand allocation. It is common in teaching simulations because it lets students see how price, quality, fit, service, reputation, technology, and lifecycle combine.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )

    _h1(doc, "4. Capacity, Materials, Production, and Defects")
    _table(
        doc,
        ["Model component", "Logic", "Managerial meaning"],
        [
            ["Effective capacity", "installed_capacity + overtime_capacity_units.", "Overtime helps this round; expansion mostly helps future rounds."],
            ["Feasible production", "Capped by effective capacity and available raw materials.", "Plans beyond constraints are reduced."],
            ["Product allocation", "If total requested production exceeds constraints, output is capped proportionally across active product requests.", "No product gets infinite priority unless the model is later extended."],
            ["Supplier cost", "Weighted material cost from offshore, balanced, and premium supplier shares, adjusted by market material cost index.", "Cheap suppliers lower cost but may raise risk/defects/lead time."],
            ["Lead time", "Weighted supplier lead time reduced by expedited order share.", "Speed costs money."],
            ["Defect rate", "Base product/archetype defect + supplier pressure + utilization stress - QC diminishing-return benefit.", "QC lowers defects, but each additional dollar helps less."],
            ["Good units", "floor(actual_production * (1 - defect_rate)).", "Defective units reduce sellable supply and increase warranty cost."],
        ],
        [1.35, 3.0, 2.2],
        fill=LIGHT_ORANGE,
        font_size=7.35,
    )

    _h1(doc, "5. Forecasting and S&OP Metrics")
    _table(
        doc,
        ["Metric", "Formula", "Interpretation"],
        [
            ["Forecast error units", "actual_demand_units - forecast_units.", "Positive means under-forecast; negative means over-forecast."],
            ["Absolute error units", "abs(actual_demand_units - forecast_units).", "Size of forecast miss regardless of direction."],
            ["Forecast bias pct", "(actual_demand_units - forecast_units) / max(forecast_units, 1).", "Direction and scale of bias."],
            ["Team WAPE", "sum(abs(error)) / max(sum(actual_demand), 1).", "Weighted absolute percentage error across products."],
            ["Forecast accuracy", "Often shown as 1 - WAPE, clipped between 0 and 1.", "Higher is better, but accuracy should be balanced with service and margin."],
            ["Forecast-production gap", "total_planned_production - total_forecast_units.", "Whether the plan makes much more or less than forecast."],
        ],
        [1.45, 2.75, 2.35],
        fill=LIGHT_BLUE,
        font_size=7.35,
    )
    _callout(
        doc,
        "Important",
        "Forecast units do not create demand. They are a planning commitment used to evaluate how well the team understood the market and aligned production, materials, inventory, and cash.",
        fill=LIGHT_RED,
        accent=RED,
    )

    _h1(doc, "6. Finance, Cash, Debt, and Working Capital")
    _table(
        doc,
        ["Financial output", "How it is computed", "Why it matters"],
        [
            ["Revenue", "sales_units * selling_price_per_unit plus any liquidation revenue.", "Top-line performance."],
            ["Procurement cost", "Raw material purchases * weighted material unit cost, plus expedite effects.", "Large orders consume cash before all demand is known."],
            ["Production/QC cost", "Actual units produced * conversion cost plus QC budget effects.", "Higher production and quality spending can raise service and lower defects, but cost cash."],
            ["Holding cost", "Ending FG/RM inventory * holding cost assumptions.", "Inventory is useful but not free."],
            ["Warranty cost", "Sales units * defect rate * warranty cost factor.", "Defects punish quality and finance."],
            ["Backlog cost", "Unmet demand carried as backlog * penalty.", "Poor service has cost and reputation effects."],
            ["Expansion cost", "Capacity expansion units * capex factor.", "Investment for future output."],
            ["Innovation investment", "Development project investment this round.", "Future products compete for current cash."],
            ["Interest expense", "Debt balance * periodic interest rate.", "Debt creates ongoing cost."],
            ["Ending cash before borrowing", "starting cash + planned borrowing + revenue - all costs.", "Can be negative even if the team reports profit."],
            ["Automatic borrowing", "If ending cash before borrowing < 0, debt increases by the deficit and cash becomes 0.", "The business survives, but debt/interest/liquidity pressure rise."],
        ],
        [1.45, 2.75, 2.35],
        fill=LIGHT_RED,
        font_size=7.2,
    )

    _h1(doc, "7. NPD Readiness and Launch Rules")
    _table(
        doc,
        ["Rule", "Model behavior", "Classroom explanation"],
        [
            ["Project creation", "Teams define project name, target segment, target tech generation, intended slot, launch timing, and modifiers.", "This is the concept/specification decision."],
            ["Locking", "After investment/testing begins, fixed project settings lock.", "Students cannot rewrite the product concept every round."],
            ["Investment progress", "Cumulative investment approaches required investment; higher tech usually requires more.", "Money speeds funding progress but has cash opportunity cost."],
            ["Testing readiness", "Testing intensity improves readiness and launch stability.", "Execution quality matters, not only budget."],
            ["Readiness threshold", "A project becomes launch-ready when funding/readiness/timing gates are satisfied.", "Launch requires resources plus maturity."],
            ["Launch now", "Team must choose launch when eligible.", "A ready project can still be delayed intentionally."],
            ["Replacement", "Launch can fill/replace an intended slot.", "New tech can improve attractiveness but may cannibalize old products."],
        ],
        [1.4, 2.75, 2.4],
        fill=LIGHT_PURPLE,
        font_size=7.35,
    )

    _h1(doc, "8. Product Lifecycle and Cannibalization")
    _table(
        doc,
        ["Feature", "Model logic", "Teaching point"],
        [
            ["Lifecycle progression", "Age maps products through launch, growth, maturity, decline.", "Products do not stay equally attractive forever."],
            ["Decline effect", "Decline lowers attractiveness and can create pressure to retire/replace.", "Old products may still sell, but less strongly."],
            ["Technology gap", "Older-than-market technology suffers more in premium and mid segments.", "Innovation timing matters most when customers care about tech."],
            ["Cannibalization", "Same-team, same-group, similar-segment newer products can transfer some demand from older products.", "Launching a new product can grow the firm or just shift demand internally."],
            ["Retirement", "Retire flag removes old products after round processing; inventory/backlog consequences are handled by the engine.", "Retirement is a portfolio cleanup decision, not just a delete button."],
        ],
        [1.45, 2.8, 2.3],
        fill=LIGHT_TEAL,
        font_size=7.35,
    )

    _h1(doc, "9. Performance Measures and Scoring")
    _table(
        doc,
        ["Measure", "What it rewards", "Potential blind spot"],
        [
            ["Profit", "Short-term economic result.", "Can reward risky debt or underinvestment."],
            ["Fill rate/service", "Ability to meet allocated demand.", "Can reward overproduction if inventory/cash ignored."],
            ["Forecast accuracy", "Planning discipline.", "A very accurate low-growth plan may miss upside."],
            ["Cash health", "Liquidity and low debt.", "May under-reward profitable expansion."],
            ["Innovation position", "Future readiness and technology advantage.", "May consume current cash."],
            ["Balanced score", "Combination of profit, service, forecast, cash, and innovation.", "Weights are instructor choices and should match learning goals."],
        ],
        [1.5, 2.6, 2.45],
        fill=LIGHT_GREEN,
        font_size=7.35,
    )
    _callout(
        doc,
        "Scoring advice",
        "Do not grade only by profit unless the lesson is deliberately short-term profit maximization. Balanced scoring makes students defend tradeoffs: innovation vs cash, volume vs margin, quality vs cost, forecast vs production, inventory vs service, growth vs liquidity.",
        fill=LIGHT_ORANGE,
        accent=ORANGE,
    )

    _h1(doc, "10. Assumptions and Tunable Parameters")
    _table(
        doc,
        ["Area", "Examples of tunable constants", "How instructors can use it"],
        [
            ["Demand environment", "Segment shares, total demand, price sensitivity, customer pickiness.", "Create premium, beginner boom, price war, recession, or boom scenarios."],
            ["Operations", "Overtime cost, capacity expansion cost, utilization stress, holding cost.", "Make capacity cheap/expensive or inventory risky/safe."],
            ["Supply", "Supplier cost multipliers, lead times, defect pressure, supply risk.", "Create cost inflation or supply shock rounds."],
            ["Quality", "QC diminishing returns, warranty factor, quality sensitivity.", "Reward/punish quality strategy."],
            ["Technology/NPD", "Required investment, readiness threshold, tech adoption, lifecycle speed.", "Make innovation more or less urgent."],
            ["Finance", "Starting cash, interest rate, debt stress thresholds.", "Teach liquidity and working capital."],
            ["Balanced score", "Weights on profit, service, forecast, cash, innovation.", "Prevent one strategy from always winning."],
        ],
        [1.35, 3.0, 2.2],
        fill=LIGHT_BLUE,
        font_size=7.35,
    )
    _h1(doc, "11. Theory Background and Classroom Interpretation")
    _table(
        doc,
        ["Theory idea", "Where it appears in Kiki", "How to explain it to students"],
        [
            ["S&OP alignment", "Forecast, production, raw materials, capacity, inventory, and cash are checked together.", "A plan is strong only when commercial demand and operating resources agree."],
            ["Resource-based competition", "Capacity, cash, reputation, product technology, and pipeline readiness carry forward.", "Teams build or lose capabilities over time; each round is not independent."],
            ["Product lifecycle theory", "Products age through launch, growth, maturity, and decline.", "A good product can become stale if the team does not refresh or replace it."],
            ["Technology adoption", "Market generation and segment tech adoption modify product attractiveness.", "New technology matters more when customers are ready to value it."],
            ["Working capital discipline", "Inventory, raw materials, backlog, debt, and interest shape cash pressure.", "Profit is not the same as liquidity; growth can break a firm."],
            ["Multi-attribute choice", "Demand is split by weighted attractiveness across price, quality, service, fit, tech, lifecycle, and reputation.", "Customers respond to bundles of value, not one isolated attribute."],
            ["Tradeoff-based strategy", "Balanced score can combine profit, service, forecast accuracy, cash, and innovation.", "There is no permanent winner; the best strategy depends on objective and environment."],
        ],
        [1.45, 2.7, 2.4],
        fill=LIGHT_TEAL,
        font_size=7.35,
    )

    out = DOC_DIR / "Kiki_Simulator_Model_and_Formula_Guide_v3_DETAILED.docx"
    doc.save(out)
    return out


def build_student_guide() -> Path:
    doc = Document()
    _style_doc(doc, landscape=True)
    _title_block(
        doc,
        "Student Guidebook",
        "Kiki Pickleball Business Simulation",
        "Detailed student-facing guide for making decisions, reading analytics, understanding products, finance, forecasts, and strategy.",
        accent=BLUE,
    )
    _callout(
        doc,
        "Your mission",
        "Run a paddle company across repeated rounds. You forecast demand, set prices, produce products, manage quality and suppliers, control cash, invest in future products, and learn from results.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    _h1(doc, "1. The Big Picture")
    _table(
        doc,
        ["What you manage", "Why it matters"],
        [
            ["Market demand", "You compete for premium, mid, and beginner customers. The market changes by round."],
            ["Operations", "Capacity, production, materials, quality, inventory, and backlog decide what you can actually sell."],
            ["Portfolio", "Product A/B/C can target different segments and age through lifecycle stages."],
            ["Development pipeline", "Future products need investment, testing, readiness, and launch timing."],
            ["Finance", "Revenue is good, but cash can disappear through procurement, production, inventory, expansion, NPD, warranty, backlog, and interest."],
            ["Learning", "Winning depends on the environment and scoring. There is no one permanent best strategy."],
        ],
        [1.55, 7.55],
        fill=LIGHT_GREEN,
        font_size=8.0,
    )

    _h1(doc, "2. What to Do Each Round")
    _numbered(
        doc,
        [
            "Read Public Market Report first.",
            "Review current team state: capacity, inventory, backlog, reputation, cash, debt, and tech position.",
            "Forecast demand by product.",
            "Choose product prices, production, QC budget, and target finished-goods inventory.",
            "Set firm-level operations: overtime, expansion, raw materials, supplier mix, expedited share, max backorders, and borrowing.",
            "Manage development projects: invest, test, launch when ready, or cancel if strategy changes.",
            "Read the Planning Analytics Preview and warnings.",
            "Save Portfolio and Pipeline Decision before the deadline.",
            "After the round, review Results Dashboard, Finance Detail, and My Team Tracker to adjust next round.",
        ],
    )

    _h1(doc, "3. Team Decision Variables")
    _table(
        doc,
        ["Variable", "Meaning", "What to think about"],
        [
            ["Forecast Units", "Your estimate of product demand.", "Forecast does not create demand; it helps align the plan and is graded after actual demand happens."],
            ["Selling Price / Unit", "Price customers pay if units sell.", "Higher price raises margin but can lower demand attractiveness."],
            ["Planned Production Units", "Units you try to make for that product.", "Production can be capped by capacity and raw materials."],
            ["QC Budget / Unit", "Quality spending per produced unit.", "Reduces defects with diminishing returns. Too little QC can create warranty/reputation damage."],
            ["Target FG Inventory", "Desired finished-goods inventory buffer.", "Protects service but creates holding/working-capital cost."],
            ["Overtime Capacity Units", "Extra short-term production capacity.", "Can help meet demand but usually increases cost/stress."],
            ["Capacity Expansion Units", "Investment in future installed capacity.", "Mostly helps later rounds and costs cash now."],
            ["Raw Material Order Qty", "Materials ordered for production.", "Too low caps production; too high ties cash/inventory."],
            ["Supplier Mix %", "Offshore, balanced, premium sourcing shares.", "Must sum to 100%; trades cost, speed, risk, and defects."],
            ["Expedited Share %", "Share of orders expedited.", "Shorter lead time but higher cost."],
            ["Max Backorder Units", "How much unmet demand can carry forward.", "Backlog can preserve demand but creates penalty and service pressure."],
            ["Planned Borrowing Amount", "Money borrowed intentionally.", "Gives cash now but increases debt and interest."],
        ],
        [1.55, 3.0, 4.55],
        fill=LIGHT_BLUE,
        font_size=7.55,
    )

    _h1(doc, "4. Product Portfolio and Product Slots")
    _callout(
        doc,
        "Before editing products",
        "Think portfolio first, then product by product. Ask which slot is supposed to win premium demand, which slot protects mid or beginner volume, and which slot may need retirement or replacement.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )
    _table(
        doc,
        ["Portfolio question", "Why it matters"],
        [
            ["Which product is our profit engine?", "High-margin products can carry profit even with lower volume."],
            ["Which product is our service/volume engine?", "Mid or beginner products may protect market share and fill rate."],
            ["Which product is aging or obsolete?", "Old products can lose attractiveness when lifecycle or technology pressure rises."],
            ["Which product should be replaced?", "A new launch can refresh the portfolio but may cannibalize old products."],
        ],
        [2.15, 6.95],
        fill=LIGHT_TEAL,
        font_size=7.75,
    )
    _figure(doc, SCREEN_DIR / "Screenshot 2026-04-30 111152.png", "Figure A. Active Product Portfolio and Forecasts: each active slot has its own product decisions.", width=9.35)
    _figure_explanation(
        doc,
        "Active product slots",
        "Product A/B/C are the products your company can sell now.",
        "For each active product, set forecast, price, production, QC, target inventory, and optionally retire after the round. Lifecycle, age, tech generation, inventory, and backlog come from prior rounds.",
        "Use the portfolio to avoid being a one-product company. Different products can target premium, mid, or beginner demand.",
        "Inactive products do not compete for demand unless launched/reactivated later.",
    )
    _table(
        doc,
        ["Field", "Meaning", "Advice"],
        [
            ["Lifecycle", "Launch, growth, maturity, or decline.", "Growth products often have stronger demand; decline products may need replacement."],
            ["Age (Rounds)", "How many rounds the product has been active.", "Older products may move toward decline."],
            ["Tech Gen", "Product technology generation.", "Higher generation helps when the market values tech."],
            ["FG Inventory", "Finished units carried into the round.", "Inventory can serve demand before new production is sold."],
            ["Backlog", "Unfilled demand carried from previous round.", "Backlog means service problems and future pressure."],
            ["Retire After Round", "Signals product should leave the portfolio after processing.", "Use for old, low-margin, low-tech, or strategically obsolete products."],
        ],
        [1.35, 3.1, 4.65],
        fill=LIGHT_TEAL,
        font_size=7.65,
    )

    _h1(doc, "5. Development Pipeline and Launch Readiness")
    _figure(doc, SCREEN_DIR / "Screenshot 2026-04-30 111205.png", "Figure B. Development Pipeline: up to two future product projects.", width=9.35)
    _figure_explanation(
        doc,
        "Development pipeline",
        "Projects are future products, not sellable products yet.",
        "Once you begin investing/testing, fixed project settings lock: project name, segment, target tech generation, intended slot, and launch timing. You can still adjust investment this round, testing intensity, launch now if ready, or cancel.",
        "Spend money and testing effort early enough so the product becomes ready before the market moves.",
        "If readiness stalls, check required funding, testing intensity, tech generation difficulty, and launch timing gates.",
        fill=LIGHT_PURPLE,
    )
    _table(
        doc,
        ["Pipeline field", "Meaning", "How it affects launch"],
        [
            ["Project Name", "Name of future product.", "Locks after project starts."],
            ["Target Segment", "Premium/mid/beginner product target.", "Affects later demand fit. Locks after project starts."],
            ["Target Tech Generation", "Technology generation you are developing.", "Higher gen can improve demand but costs more and may be harder to ready. Locks after project starts."],
            ["Intended Product Slot", "A/B/C slot the project should fill or replace.", "Defines replacement plan. Locks after project starts."],
            ["Planned Launch Round", "When you planned to launch.", "Past planned launch should not block launching later if readiness is reached."],
            ["Investment This Round", "Money spent now.", "Raises cumulative investment and can speed progress but consumes cash."],
            ["Testing Intensity", "Effort to stabilize/validate product.", "Improves readiness and may reduce launch risk."],
            ["Launch Now If Ready", "Team decision to launch eligible project.", "Launch occurs only if gates are satisfied."],
            ["Cancel Project", "Stop the project.", "Use if strategy changes or cash is too tight."],
        ],
        [1.45, 3.1, 4.55],
        fill=LIGHT_PURPLE,
        font_size=7.45,
    )

    _h1(doc, "6. Planning Analytics Preview")
    _figure(doc, SCREEN_DIR / "Screenshot 2026-04-30 111223.png", "Figure C. Planning Analytics Preview: decision-quality checkpoint before saving.", width=9.35)
    _figure_explanation(
        doc,
        "Planning preview",
        "This is your pre-submit dashboard.",
        "It estimates total forecast, capacity, production, utilization, raw material sufficiency, defect rate, margin/unit, working capital, projected cash, borrowing need, pipeline status, tech position, decline products, and cannibalization.",
        "Do not ignore yellow warnings. They usually tell you where the plan is incoherent.",
        "Preview values are estimates, not guaranteed results, because competitors and final demand allocation still matter.",
        fill=LIGHT_ORANGE,
    )
    _table(
        doc,
        ["Preview metric", "Meaning", "If it looks bad..."],
        [
            ["Total Forecast Units", "Your total expected demand across active products.", "Recheck product forecasts and market report."],
            ["Effective Firm Capacity", "Installed capacity plus overtime.", "Add overtime/expansion or reduce production."],
            ["Total Planned Production", "All product production combined.", "Compare to forecast and capacity."],
            ["Forecast-Production Gap", "Production minus forecast.", "Large positive gap risks inventory; large negative gap risks stockout."],
            ["Planned Utilization %", "How much capacity you plan to use.", "Very high utilization can raise stress/defects."],
            ["Raw Material Available", "Beginning RM plus new material position used in planning.", "Order more or reduce production if insufficient."],
            ["Projected Margin / Unit", "Selling price minus estimated unit cost.", "If negative, selling more can lose more money."],
            ["Projected Cash Before Borrowing", "Cash estimate before emergency borrowing.", "If negative, expect debt."],
            ["Likely Borrowing Need", "Estimated automatic borrowing need.", "Reduce spending or intentionally plan borrowing."],
            ["Likely Cannibalization", "How much your products may steal from each other.", "Avoid launching products that only replace your own demand unless strategically needed."],
        ],
        [1.7, 3.1, 4.3],
        fill=LIGHT_ORANGE,
        font_size=7.35,
    )

    _h1(doc, "7. How Profit, Cost, Cash, and Debt Happen")
    _table(
        doc,
        ["Term", "What it means in the game"],
        [
            ["Revenue", "Units sold times selling price. More demand only helps if you have sellable units."],
            ["Procurement cost", "Cost of raw material orders, affected by supplier mix, material index, and expedite decisions."],
            ["Production cost", "Cost of making actual units, including conversion/overtime/QC logic."],
            ["Holding cost", "Cost of leftover finished goods and raw materials."],
            ["Warranty cost", "Cost caused by defects on units sold."],
            ["Backlog cost", "Penalty for demand you could not serve and carried forward."],
            ["Expansion cost", "Cash spent to add future capacity."],
            ["Innovation investment", "Cash spent on development pipeline."],
            ["Interest expense", "Cost of debt. Debt is not free money."],
            ["Profit", "Revenue minus costs."],
            ["Cash", "Money available after revenue, spending, borrowing, and debt effects. You can be profitable but cash-constrained."],
            ["Automatic borrowing", "If cash would go below zero, the simulator adds short-term debt so your cash does not become negative."],
        ],
        [1.7, 7.4],
        fill=LIGHT_RED,
        font_size=7.7,
    )

    _h1(doc, "8. Big-Picture Figures")
    figure_data = [
        ("Figure_1_The_Round_Loop.png", "The Round Loop", "Shows the repeated cycle: read market, submit decisions, engine calculates results, learn and adjust.", "Inventory, backlog, capacity, reputation, cash, debt, product age, and pipeline readiness carry forward.", "Use after every round: what did we learn and what must change next?"),
        ("Figure_2_What_Students_Control.png", "What Students Control", "Maps six decision areas: demand, operations, supply, portfolio, pipeline, finance.", "No single control wins alone; the plan must fit together.", "Before saving, ask whether each decision area supports the same strategy."),
        ("Figure_3_Portfolio_and_Pipeline.png", "Portfolio and Pipeline", "Separates active products from future development projects and launch gates.", "Active products earn now; pipeline projects prepare future competitiveness.", "Use this to decide whether to exploit current products or invest in future products."),
        ("Figure_4_Demand_Allocation_Infographic.png", "Demand Allocation", "Explains market segments, attractiveness scores, and proportional demand share.", "You do not receive fixed demand; you earn relative demand against competitors.", "Improve attractiveness through price fit, quality, service, technology, lifecycle, reputation, and segment alignment."),
        ("Figure_5_Operations_and_Cost_Engine.png", "Operations and Cost Engine", "Shows capacity/material constraints, defects, good units, sales service, and cost buckets.", "Demand is useless without capacity/materials/good units.", "Use it to explain why high demand can still become backlog, lost sales, defects, or cost pressure."),
        ("Figure_6_Forecasting_and_SOP.png", "Forecasting and S&OP", "Shows forecast, planning response, and after-round metrics.", "Forecasting disciplines the plan; it does not create demand.", "Use forecast accuracy to learn whether your team understood the market."),
        ("Figure_7_Cash_and_Debt_Pressure.png", "Cash and Debt Pressure", "Shows cash bridge from starting cash to ending cash and borrowing rule.", "Profit and cash are different; debt creates interest and liquidity stress.", "Use this when deciding expansion/NPD/inventory spending."),
        ("Figure_8_Strategy_Archetypes.png", "Strategy Archetypes", "Compares six common strategies and their strengths/risks.", "There is no permanent winner; market conditions and scoring weights matter.", "Pick a strategy deliberately, then adapt when the environment changes."),
    ]
    for filename, title, purpose, how_to_read, use in figure_data:
        _figure(doc, FIG_DIR / filename, f"{title}. Classroom infographic for Kiki Pickleball Business Simulation.", width=9.25)
        _figure_explanation(
            doc,
            title,
            purpose,
            how_to_read,
            use,
            "Use the figure as a conversation tool, not as a rule sheet. Your decisions still need numbers that fit the current market.",
            fill=LIGHT_TEAL,
        )

    _h1(doc, "9. Strategy Advice")
    _table(
        doc,
        ["Strategy", "Good when...", "Main risk"],
        [
            ["Cash Conservative", "Supply/cost shock, recession, cash crunch, inventory risk.", "May miss growth or innovation upside."],
            ["Balanced S&OP", "Forecast volatility, stable markets, balanced scoring.", "May be too cautious in boom markets."],
            ["Premium Quality", "Premium market, picky customers, high quality sensitivity.", "Higher QC/cost can hurt if customers become price sensitive."],
            ["Innovation Leap", "Tech shift, premium tech adoption, fast lifecycle.", "Pipeline investment and launch risk can drain cash."],
            ["Aggressive Growth", "Demand boom, capacity-constrained opportunity.", "Debt, inventory, overtime, and cost pressure."],
            ["Low-Cost Volume", "Beginner boom, price war, price-sensitive market.", "Weak margins and quality pressure."],
        ],
        [1.6, 4.0, 3.5],
        fill=LIGHT_GREEN,
        font_size=7.55,
    )
    _table(
        doc,
        ["After each round, ask...", "Why it matters for the next decision"],
        [
            ["Did we win demand but lose money?", "Check margin, cost buckets, defects, and inventory/working-capital pressure."],
            ["Did we forecast accurately?", "Use forecast error to adjust next product forecasts before changing production."],
            ["Did we stock out or build inventory?", "Match production, raw materials, and target inventory to realistic demand."],
            ["Did cash fall or debt rise?", "Separate profitable growth from liquidity stress."],
            ["Did any product age into decline?", "Consider retirement, replacement, or a targeted NPD launch."],
            ["Did the market generation move?", "Re-check whether the portfolio technology still fits customer expectations."],
        ],
        [2.2, 6.9],
        fill=LIGHT_ORANGE,
        font_size=7.65,
    )
    _table(
        doc,
        ["Final save checklist", "Complete before deadline"],
        [
            ["Market report read", "[ ]"],
            ["Supplier mix sums to 100%", "[ ]"],
            ["Every active product has a forecast", "[ ]"],
            ["Production is feasible with capacity and raw materials", "[ ]"],
            ["Price is above projected variable cost unless intentionally strategic", "[ ]"],
            ["Cash/borrowing impact has been checked", "[ ]"],
            ["Pipeline investment/testing choices are intentional", "[ ]"],
            ["Save Portfolio and Pipeline Decision clicked", "[ ]"],
        ],
        [7.7, 1.4],
        fill=LIGHT_BLUE,
        font_size=7.65,
    )
    _callout(
        doc,
        "Final student reminder",
        "A strong plan is coherent: forecasts, prices, production, suppliers, QC, inventory, cash, and pipeline investments all support the same strategy for the current market.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    out = DOC_DIR / "Kiki_Student_Guidebook_with_Figures_v2_DETAILED.docx"
    doc.save(out)
    return out


def build_offline_experiment_guide() -> Path:
    doc = Document()
    _style_doc(doc, landscape=True)
    _title_block(
        doc,
        "Offline 20-Environment Experiment Guide",
        "Instructor and researcher guide for running the simulation-of-simulation experiment.",
        "Audience: instructors, TAs, researchers, and developers calibrating Kiki outside the live classroom database.",
        accent=PURPLE,
    )
    _callout(
        doc,
        "Purpose",
        "This offline runner simulates the simulator under 20 different named market environments. It is for calibration, teaching examples, and research-style sensitivity analysis. It does not change the classroom SQLite database.",
        fill=LIGHT_PURPLE,
        accent=PURPLE,
    )
    _callout(
        doc,
        "Important",
        "Run this from PowerShell in the project folder. The classroom app can be closed; this script writes only to simulation_outputs and produces Excel/CSV files.",
        fill=LIGHT_RED,
        accent=RED,
    )

    _h1(
        doc,
        "1. What This Experiment Is",
        "The live simulator is the classroom game. The offline experiment is a separate analysis tool that repeatedly runs strategy archetypes through controlled market worlds so the instructor can see whether different environments reward different strategies.",
    )
    _table(
        doc,
        ["Item", "Meaning", "Instructor use"],
        [
            ["Live classroom app", "Students log in, submit decisions, and the instructor runs rounds.", "Use for class play and graded simulation sessions."],
            ["Offline experiment runner", "A script runs synthetic teams through 20 market environments.", "Use before class to calibrate, create examples, or support research analysis."],
            ["No database writes", "The script does not edit users, student submissions, or classroom results.", "Safe to run without damaging a live class."],
            ["Deterministic seed", "The same settings produce repeatable output.", "Useful for comparing revisions across simulator versions."],
        ],
        [1.7, 4.0, 4.0],
        fill=LIGHT_PURPLE,
        font_size=7.8,
    )

    _h1(doc, "2. Files and Commands")
    _table(
        doc,
        ["File", "Role"],
        [
            ["scripts/run_20_environment_simulations.py", "Main offline experiment orchestrator. It loops through 20 environments, runs all strategies, and exports Excel/CSV outputs."],
            ["scripts/run_strategy_simulation.py", "Supporting strategy simulation runner used by the 20-environment script."],
            ["simulation_outputs/", "Output folder. Each run creates a timestamped subfolder so old experiments are preserved."],
            ["requirements.txt", "Python dependencies for the app and offline experiment, including pandas/openpyxl style workbook support."],
        ],
        [3.0, 6.6],
        fill=LIGHT_BLUE,
        font_size=7.8,
    )
    _table(
        doc,
        ["PowerShell command", "What it does"],
        [
            [r"cd C:\Users\mk156\Documents\Codex\2026-04-21-build-a-minimal-local-simulator-app", "Move into the project folder. This avoids path/import errors."],
            [r"python scripts\run_20_environment_simulations.py --teams 6 --rounds 12", "Run baseline mode with six strategies and 12 rounds."],
            [r"python scripts\run_20_environment_simulations.py --teams 6 --rounds 12 --calibrated", "Run calibrated mode. Scenario-specific logic is stronger so different environments can reward different strategies."],
        ],
        [4.8, 4.8],
        fill=LIGHT_GREEN,
        font_size=7.35,
    )

    _h1(
        doc,
        "3. Experiment Design",
        "Every environment uses the same six strategy archetypes. The environment changes the market and operating pressure around them.",
    )
    _table(
        doc,
        ["Design choice", "Setting"],
        [
            ["Number of environments", "20 named market environments."],
            ["Teams per environment", "Default six synthetic teams, one per strategy archetype."],
            ["Rounds per environment", "Default 12 rounds. You can change this with --rounds."],
            ["Strategies held constant", "Cash Conservative, Balanced S&OP, Premium Quality, Innovation Leap, Aggressive Growth, Low-Cost Volume."],
            ["Environment changes", "Demand, segment shares, price/quality/service sensitivity, technology pressure, supply risk, material cost, forecast volatility, cash stress, inventory risk, and lifecycle pressure."],
            ["Output logic", "Each environment gets a detail sheet, plus cross-scenario summary sheets and CSV exports."],
        ],
        [2.2, 7.4],
        fill=LIGHT_TEAL,
        font_size=7.55,
    )

    _h1(doc, "4. The 20 Market Environments")
    _table(
        doc,
        ["#", "Environment", "What it tests", "#", "Environment", "What it tests"],
        [
            ["1", "Baseline", "Balanced demand, normal costs, moderate tech pressure.", "11", "Fast Lifecycle", "Products age faster; pipeline timing matters."],
            ["2", "Picky Customers", "Quality, reputation, and service matter more.", "12", "Warranty Sensitive", "Defects are more expensive and visible."],
            ["3", "Price Sensitive", "Customers react strongly to price differences.", "13", "Price War", "Affordability matters, but margins are at risk."],
            ["4", "Tech Shift", "Newer technology gains demand advantage.", "14", "Inventory Risk", "Holding excess stock becomes dangerous."],
            ["5", "Beginner Boom", "Beginner segment expands with price pressure.", "15", "Capacity Constraint", "Capacity bottlenecks shape production choices."],
            ["6", "Premium Market", "Premium demand and tech adoption increase.", "16", "Forecast Volatility", "Demand is harder to predict accurately."],
            ["7", "Demand Recession", "Demand falls; overproduction and debt hurt.", "17", "Viral Demand", "Demand spikes; fast response can win."],
            ["8", "Demand Boom", "Demand grows; scaling creates opportunity and risk.", "18", "Cash Crunch", "Liquidity and debt discipline matter more."],
            ["9", "Supply Shock", "Supply risk and input instability increase.", "19", "Premium Expansion", "Premium customers expect stronger tech fit."],
            ["10", "Cost Inflation", "Input costs rise and squeeze contribution margin.", "20", "Unstable Market", "Mixed shocks reward resilience and balance."],
        ],
        [0.3, 1.35, 3.05, 0.3, 1.35, 3.05],
        fill=LIGHT_PURPLE,
        font_size=6.9,
    )

    _h1(
        doc,
        "5. Strategy Archetypes Tested",
        "These six archetypes run in every environment, which lets you compare how the same strategy performs when the business world changes.",
    )
    _table(
        doc,
        ["Strategy", "Operating archetype used", "Best-fit environment", "Typical strength", "Typical weakness"],
        [
            ["Cash Conservative", "Community / Club", "Supply Shock, Cash Crunch, Inventory Risk, Recession", "Protects liquidity, avoids overexpansion, limits downside.", "May miss boom, innovation, and premium upside."],
            ["Balanced S&OP", "Engineering Challenger", "Forecast Volatility, Baseline, mixed environments", "Matches forecast, production, materials, and capacity coherently.", "Can be too cautious when fast growth is needed."],
            ["Premium Quality", "Premium Quality", "Picky Customers, Premium Market, Warranty Sensitive", "Wins with quality, reputation, QC, and premium fit.", "Higher cost can hurt in price-sensitive worlds."],
            ["Innovation Leap", "Engineering Challenger", "Tech Shift, Fast Lifecycle, Premium Expansion", "Creates upside through new generation products and launch timing.", "Pipeline investment and launch risk can drain cash."],
            ["Aggressive Growth", "Sponsorship Growth", "Demand Boom, Viral Demand, Capacity opportunity", "Uses capacity and production to capture rapid demand.", "Can create debt, inventory, and execution stress."],
            ["Low-Cost Volume", "Low-Cost Volume", "Beginner Boom, Price Sensitive, Price War", "Competes on affordability and unit volume.", "Weak margins and quality pressure can erase sales gains."],
        ],
        [1.25, 1.45, 2.0, 2.45, 2.45],
        fill=LIGHT_GREEN,
        font_size=6.85,
    )
    _callout(
        doc,
        "NPD modeling note",
        "The offline runner also gives each strategy a different development-project profile. Innovation Leap targets newer technology with higher demand-fit ambition; Premium Quality targets a cleaner premium product; Low-Cost Volume develops a value refresh; Cash Conservative invests slowly.",
        fill=LIGHT_PURPLE,
        accent=PURPLE,
    )

    _h1(doc, "6. Environment Control Variables")
    _table(
        doc,
        ["Variable", "Plain-English meaning", "What changing it usually teaches"],
        [
            ["demand_multiplier", "Scales total market demand up or down.", "Boom vs recession; capacity and cash discipline."],
            ["demand_growth_delta", "Changes how demand evolves over rounds.", "Growth planning and over/under-scaling."],
            ["premium/mid/beginner share shifts", "Changes which customer segments are largest.", "Portfolio positioning and segment fit."],
            ["price_sensitivity", "Makes customers react more strongly to price.", "Low-cost strategy, margin protection, price war behavior."],
            ["quality_sensitivity / customer_pickiness", "Makes QC, defects, service, and reputation matter more.", "Premium strategy and warranty/service tradeoffs."],
            ["service sensitivity", "Rewards availability and fill rate more heavily.", "Inventory, backlog, capacity, and S&OP discipline."],
            ["technology_shift_delta / market generation", "Moves the market toward newer technology.", "NPD timing, product obsolescence, innovation advantage."],
            ["material_cost_index_delta", "Raises or lowers input cost.", "Supplier mix, pricing, and contribution margin."],
            ["supply_risk_override", "Changes supplier reliability/defect pressure.", "Sourcing resilience and quality risk."],
            ["forecast_volatility", "Makes demand harder to predict.", "Forecast accuracy, S&OP discipline, and planning buffers."],
            ["cash_stress_factor", "Raises pressure from debt, liquidity, and financing.", "Working capital, borrowing, and growth discipline."],
        ],
        [2.1, 3.55, 3.95],
        fill=LIGHT_ORANGE,
        font_size=6.95,
    )

    _h1(doc, "7. Baseline Mode vs Calibrated Mode")
    _table(
        doc,
        ["Mode", "Command flag", "Best use", "Interpretation"],
        [
            ["Baseline", "No --calibrated flag", "See what the current model naturally rewards.", "If one strategy dominates, the game may need stronger scenario-specific incentives."],
            ["Calibrated", "--calibrated", "Create better classroom variety across environments.", "Scenario-specific pressures are strengthened so Tech Shift, Premium Market, Price War, Cash Crunch, and Forecast Volatility produce different winners more often."],
        ],
        [1.3, 1.8, 3.0, 3.5],
        fill=LIGHT_BLUE,
        font_size=7.3,
    )
    _callout(
        doc,
        "Calibration warning",
        "If one strategy wins more than 60% of profit or balanced scenarios, the script prints a warning. That does not mean the script failed; it means the model is telling you one strategy is too strong under the current calibration.",
        fill=LIGHT_RED,
        accent=RED,
    )

    _h1(doc, "8. Workbook and CSV Outputs")
    _table(
        doc,
        ["Output", "What it contains", "How instructor should use it"],
        [
            ["pickleball_20_environment_simulations.xlsx", "One workbook with summary sheets plus one sheet per environment.", "Main file for teaching, inspection, and calibration."],
            ["Scenario_Assumptions sheet", "One row per environment with all parameter values.", "Check what actually changed between environments."],
            ["All_Scenario_Winners sheet", "Profit, balanced, service, forecast, cash, and innovation winners.", "Show students that winners depend on objective and environment."],
            ["Strategy_Summary sheet", "Win counts and average performance by strategy.", "See whether the strategy set is balanced enough."],
            ["Teaching_Insights sheet", "Scenario interpretation, best strategy, lesson, and discussion question.", "Use directly during class debrief."],
            ["Scenario detail sheets", "Assumptions, final team results, round-by-round details, and winner summary.", "Inspect exactly how a strategy won or failed."],
            ["CSV files", "Same major tables exported separately.", "Use for statistics, R/Python analysis, or paper appendices."],
        ],
        [2.25, 3.55, 3.85],
        fill=LIGHT_TEAL,
        font_size=7.0,
    )
    _table(
        doc,
        ["CSV file", "Use"],
        [
            ["scenario_assumptions.csv", "Environment parameter table."],
            ["all_scenario_winners.csv", "Winner by metric for each scenario."],
            ["strategy_summary.csv", "Cross-scenario strategy aggregation."],
            ["teaching_insights.csv", "Ready-to-use classroom interpretation prompts."],
            ["round_level_results.csv", "One row per team per round for detailed time-series analysis."],
            ["final_team_results.csv", "Final team outcomes across all environments."],
        ],
        [3.2, 6.4],
        fill=LIGHT_GRAY,
        font_size=7.35,
    )

    _h1(doc, "9. How to Interpret Winners")
    _table(
        doc,
        ["Winner type", "Metric idea", "Teaching interpretation"],
        [
            ["Profit winner", "Highest final profit.", "Best financial upside in that environment, but may hide debt or service risk."],
            ["Balanced winner", "Composite score across profit, service, forecast discipline, cash, and innovation.", "Best overall managerial performance."],
            ["Service winner", "Best fill/service performance.", "Best at meeting demand and avoiding stockouts/backlog."],
            ["Forecast winner", "Best forecast accuracy / lowest planning error.", "Best S&OP discipline, not necessarily highest sales."],
            ["Cash winner", "Strongest cash/debt position.", "Best liquidity discipline and working-capital control."],
            ["Innovation winner", "Best innovation/pipeline/technology position.", "Best future-readiness, which may or may not pay off immediately."],
        ],
        [1.65, 2.8, 5.15],
        fill=LIGHT_GREEN,
        font_size=7.3,
    )
    _callout(
        doc,
        "Instructor framing",
        "Do not tell students there is one permanent best strategy. Use the output to ask: best for what objective, in what market, with what risk?",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    _h1(
        doc,
        "10. Classroom and Research Workflow",
        "Use the offline outputs before class to calibrate scenarios, and after class to connect student decisions to benchmark strategy patterns.",
    )
    _table(
        doc,
        ["Step", "Action", "Why it matters"],
        [
            ["1", "Run calibrated 20-environment experiment before class.", "Confirms the simulator creates strategy variety."],
            ["2", "Open Strategy_Summary and All_Scenario_Winners.", "Identify which strategies win under which conditions."],
            ["3", "Pick two contrasting environments.", "For example, Cash Crunch vs Tech Shift or Price War vs Premium Market."],
            ["4", "Use the output as a teaching preview.", "Show that the same strategy can be strong in one world and weak in another."],
            ["5", "After live class rounds, compare student outcomes to offline archetypes.", "Helps students diagnose whether their behavior resembled conservative, growth, premium, or innovation logic."],
            ["6", "For research, archive the workbook and CSVs with version/date.", "Keeps calibration evidence reproducible for paper development."],
        ],
        [0.45, 4.15, 5.0],
        fill=LIGHT_PURPLE,
        font_size=7.35,
    )

    _h1(doc, "11. Troubleshooting")
    _table(
        doc,
        ["Problem", "Likely cause", "Fix"],
        [
            ["PowerShell says file not found", "Command was run outside the project folder.", "Run cd into the project folder first."],
            ["ModuleNotFoundError", "Dependencies are not installed in the active Python environment.", "Run pip install -r requirements.txt or use the same Python environment as the app."],
            ["No Excel file appears", "Script failed before writing outputs or user looked in wrong folder.", "Check terminal output for the exact Output folder path."],
            ["Excel workbook will not open/save", "The previous workbook is open in Excel.", "Close Excel or rerun to create a new timestamped folder."],
            ["One strategy wins everything", "Scenario-specific pressure is too weak or scoring weights reward one behavior.", "Use --calibrated and review the calibration warning."],
            ["Instructor worries about live class data", "Confusion between offline outputs and classroom SQLite.", "This runner writes to simulation_outputs only, not simulator.db."],
        ],
        [2.2, 3.4, 4.0],
        fill=LIGHT_RED,
        font_size=7.15,
    )

    _h1(doc, "12. Quick Start Checklist")
    _table(
        doc,
        ["Done", "Task"],
        [
            ["[ ]", "Open PowerShell."],
            ["[ ]", r"cd C:\Users\mk156\Documents\Codex\2026-04-21-build-a-minimal-local-simulator-app"],
            ["[ ]", r"Run python scripts\run_20_environment_simulations.py --teams 6 --rounds 12 --calibrated"],
            ["[ ]", "Wait for progress messages [1/20] through [20/20]."],
            ["[ ]", "Open the printed Excel workbook path."],
            ["[ ]", "Review All_Scenario_Winners and Strategy_Summary."],
            ["[ ]", "Use Teaching_Insights to prepare debrief prompts."],
            ["[ ]", "Archive the timestamped output folder if using it for research or calibration evidence."],
        ],
        [0.65, 8.95],
        fill=LIGHT_ORANGE,
        font_size=7.35,
    )

    out = DOC_DIR / "Kiki_Offline_20_Environment_Experiment_Guide_v1_DETAILED.docx"
    doc.save(out)
    return out


def _latest_calibrated_experiment_output() -> Path | None:
    """Return the newest calibrated 20-environment output folder, if present."""
    if not SIMULATION_OUTPUTS_DIR.exists():
        return None
    candidates = [
        item
        for item in SIMULATION_OUTPUTS_DIR.iterdir()
        if item.is_dir() and item.name.startswith("twenty_environment_calibrated_")
    ]
    if not candidates:
        return None
    return max(candidates, key=lambda item: item.stat().st_mtime)


def main() -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [
        build_instructor_manual(),
        build_model_guide(),
        build_student_guide(),
        build_offline_experiment_guide(),
    ]
    zip_path = DOC_DIR / "Kiki_Documentation_Pack_v4_DETAILED.zip"
    if zip_path.exists():
        zip_path.unlink()
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as archive:
        for output in outputs:
            archive.write(output, output.name)
        for script in (EXPERIMENT_SCRIPT, STRATEGY_RUNNER_SCRIPT, REQUIREMENTS_FILE):
            if script.exists():
                archive.write(script, f"experiment_code/{script.name}")
        latest_experiment_output = _latest_calibrated_experiment_output()
        if latest_experiment_output is not None:
            for file_path in sorted(latest_experiment_output.iterdir()):
                if file_path.is_file() and file_path.suffix.lower() in {".xlsx", ".csv"}:
                    archive.write(
                        file_path,
                        f"experiment_outputs/{latest_experiment_output.name}/{file_path.name}",
                    )
    outputs.append(zip_path)
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
